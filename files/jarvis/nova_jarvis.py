"""
NOVA JARVIS - Asistente todopoderoso con IA
============================================
Un asistente estilo Iron Man con SUPERPODERES:

  - Chat con memoria persistente (hechos + conversaciones)
  - Modo voz continuo (escucha y responde hablando)
  - Herramientas ejecutables por la IA:
      * abrir_web(url), buscar_web(query)
      * abrir_app(nombre), listar_carpeta(ruta)
      * leer_archivo, escribir_archivo
      * ejecutar_comando (con confirmacion)
      * captura_pantalla, analizar_imagen
      * generar_imagen (Pollinations)
      * decir_hora, decir_clima
      * recordar_hecho, olvidar_hecho
      * calcular_expresion
      * traducir, resumir
      * crear_recordatorio, listar_recordatorios
  - Esfera animada con 4 estados (idle/listening/thinking/speaking)
  - Recordatorios con notificaciones
  - Base de conocimiento personal ampliable
  - Comandos rapidos con "/"
"""
import os, sys, json, threading, time, re, subprocess, webbrowser
import urllib.request, urllib.parse
import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog
from datetime import datetime, timedelta

APP_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# VERSION DEL MODULO (para el sistema de actualizaciones)
# ============================================================
VERSION = "2.0.0"
MODULO_ID = "jarvis"
PARENT = os.path.dirname(APP_DIR)
CONFIG = os.path.join(PARENT, "config.json")
MEMORIA = os.path.join(APP_DIR, "jarvis_memoria.json")
RECORDATORIOS = os.path.join(APP_DIR, "recordatorios.json")
RUTINAS = os.path.join(APP_DIR, "rutinas.json")
CAPTURAS = os.path.join(APP_DIR, "capturas")
os.makedirs(CAPTURAS, exist_ok=True)

sys.path.insert(0, APP_DIR)
from nova_ia_lib import llamar_llm, dialogo_config
from nova_brain import Brain, extraer_hechos
try:
    from nova_voz import VozProfesional
    VOZ_PRO_OK = True
except ImportError:
    VOZ_PRO_OK = False
    VozProfesional = None
try:
    from nova_rag import NovaRAG
    RAG_OK = True
except ImportError:
    RAG_OK = False
    NovaRAG = None
try:
    from nova_personalidades import PERSONALIDADES, obtener_prompt, obtener_voz, lista_personalidades
    PERS_OK = True
except ImportError:
    PERS_OK = False
    PERSONALIDADES = {}
try:
    import nova_musica
    MUSICA_OK = True
except ImportError:
    MUSICA_OK = False
    nova_musica = None
try:
    from nova_updater import Updater, mostrar_dialogo_update
    UPDATER_OK = True
except ImportError:
    UPDATER_OK = False
    Updater = None
    mostrar_dialogo_update = None

# Dependencias opcionales
try:
    import edge_tts, asyncio
    EDGE_OK = True
except ImportError:
    EDGE_OK = False

try:
    import speech_recognition as sr
    SR_OK = True
except ImportError:
    SR_OK = False

try:
    from PIL import ImageGrab, Image
    PIL_OK = True
except ImportError:
    PIL_OK = False

try:
    import pygame
    # NO inicializar el mixer aqui - se hace bajo demanda
    PYGAME_OK = True
except Exception:
    PYGAME_OK = False

_pygame_init = False
def _init_pygame():
    """Inicializa pygame mixer solo cuando se necesite (lazy)."""
    global _pygame_init
    if _pygame_init or not PYGAME_OK:
        return _pygame_init
    try:
        pygame.mixer.init()
        _pygame_init = True
    except Exception as e:
        print(f"pygame mixer no disponible: {e}")
    return _pygame_init

# ============================================================
# PALETA DE COLORES
# ============================================================
C = {
    "bg":       "#050a1a",
    "bg2":      "#0d1729",
    "bg3":      "#1a2540",
    "fg":       "#e0e7ff",
    "fg2":      "#94a3b8",
    "user":     "#60a5fa",
    "jarvis":   "#a78bfa",
    "sys":      "#94a3b8",
    "tool":     "#22c55e",
    "err":      "#ef4444",
    "warn":     "#fbbf24",
    "ok":       "#22c55e",
    "orb_idle":     "#3b82f6",
    "orb_listen":   "#22c55e",
    "orb_think":    "#fbbf24",
    "orb_speak":    "#a78bfa",
    "orb_tool":     "#ec4899",
}

VOCES = {
    "Elvira (Espana)":  "es-ES-ElviraNeural",
    "Alvaro (Espana)":  "es-ES-AlvaroNeural",
    "Ximena (Espana)":  "es-ES-XimenaNeural",
    "Dalia (Mexico)":   "es-MX-DaliaNeural",
    "Jorge (Mexico)":   "es-MX-JorgeNeural",
    "Elena (Argentina)":"es-AR-ElenaNeural",
    "Salome (Colombia)":"es-CO-SalomeNeural",
}

# ============================================================
# UTILIDADES DE ARCHIVO
# ============================================================
def cargar_json(ruta, default=None):
    try:
        with open(ruta, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default if default is not None else {}

def guardar_json(ruta, data):
    try:
        with open(ruta, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error guardando {ruta}: {e}")

# ============================================================
# SISTEMA DE HERRAMIENTAS (TOOL CALLING)
# ============================================================
# Todas las herramientas devuelven (resultado_str, exito_bool)

class Herramientas:
    """Coleccion de herramientas que la IA puede invocar."""
    def __init__(self, jarvis):
        self.jarvis = jarvis

    # === WEB ===
    def abrir_web(self, url):
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        try:
            webbrowser.open(url)
            return f"Abierto en el navegador: {url}", True
        except Exception as e:
            return f"Error: {e}", False

    def buscar_web(self, query):
        url = f"https://www.google.com/search?q={urllib.parse.quote(query)}"
        try:
            webbrowser.open(url)
            return f"Buscando '{query}' en Google", True
        except Exception as e:
            return f"Error: {e}", False

    def buscar_youtube(self, query):
        url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(query)}"
        try:
            webbrowser.open(url)
            return f"Buscando '{query}' en YouTube", True
        except Exception as e:
            return f"Error: {e}", False

    def buscar_wikipedia(self, query):
        try:
            # API REST de Wikipedia (sin librerias externas)
            q = urllib.parse.quote(query)
            url = f"https://es.wikipedia.org/api/rest_v1/page/summary/{q}"
            req = urllib.request.Request(url,
                headers={"User-Agent": "NovaJarvis/2.0"})
            with urllib.request.urlopen(req, timeout=10) as r:
                data = json.loads(r.read())
            resumen = data.get("extract", "Sin resumen disponible")
            return resumen[:800], True
        except Exception as e:
            return f"No encontrado en Wikipedia: {e}", False

    # === SISTEMA ===
    def abrir_app(self, nombre):
        """Abre una aplicacion del sistema."""
        try:
            if sys.platform == "win32":
                # Windows: usa start
                subprocess.Popen(f'start "" "{nombre}"', shell=True)
            else:
                subprocess.Popen([nombre])
            return f"Intentando abrir: {nombre}", True
        except Exception as e:
            return f"Error abriendo {nombre}: {e}", False

    def listar_carpeta(self, ruta):
        try:
            ruta = os.path.expanduser(ruta) if ruta else "."
            if not os.path.isdir(ruta):
                return f"No es una carpeta valida: {ruta}", False
            items = os.listdir(ruta)
            items.sort()
            lista = "\n".join(f"  {'[D]' if os.path.isdir(os.path.join(ruta, x)) else '[F]'} {x}"
                                for x in items[:50])
            n = len(items)
            return f"Carpeta {ruta} ({n} items):\n{lista}", True
        except Exception as e:
            return f"Error: {e}", False

    def leer_archivo(self, ruta):
        try:
            ruta = os.path.expanduser(ruta)
            with open(ruta, "r", encoding="utf-8", errors="ignore") as f:
                contenido = f.read()
            if len(contenido) > 3000:
                return contenido[:3000] + "\n\n[...truncado, archivo muy grande]", True
            return contenido, True
        except Exception as e:
            return f"Error leyendo: {e}", False

    def escribir_archivo(self, ruta, contenido):
        try:
            ruta = os.path.expanduser(ruta)
            os.makedirs(os.path.dirname(os.path.abspath(ruta)), exist_ok=True)
            with open(ruta, "w", encoding="utf-8") as f:
                f.write(contenido)
            return f"Archivo guardado: {ruta} ({len(contenido)} chars)", True
        except Exception as e:
            return f"Error escribiendo: {e}", False

    def ejecutar_comando(self, cmd):
        # Confirmacion importante
        if not messagebox.askyesno("Ejecutar comando?",
            f"Jarvis quiere ejecutar:\n\n  {cmd}\n\n"
            "PERMITIR? (solo Si si confias en el comando)"):
            return "Comando cancelado por el usuario", False
        try:
            result = subprocess.run(cmd, shell=True, capture_output=True,
                                     text=True, timeout=30)
            out = (result.stdout or "") + (result.stderr or "")
            return out[:2000] if out else "Comando ejecutado (sin salida)", True
        except subprocess.TimeoutExpired:
            return "Timeout: el comando tardo mas de 30 segundos", False
        except Exception as e:
            return f"Error: {e}", False

    # === PANTALLA/IMAGEN ===
    def captura_pantalla(self):
        if not PIL_OK:
            return "Necesitas Pillow: pip install Pillow", False
        try:
            img = ImageGrab.grab()
            ruta = os.path.join(CAPTURAS,
                f"captura_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
            img.save(ruta)
            return f"Captura guardada: {ruta}", True
        except Exception as e:
            return f"Error: {e}", False

    def generar_imagen(self, prompt):
        try:
            slug = urllib.parse.quote(prompt)
            url = (f"https://image.pollinations.ai/prompt/{slug}"
                   f"?width=1024&height=1024&nologo=true")
            ruta = os.path.join(CAPTURAS,
                f"img_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
            req = urllib.request.Request(url,
                headers={"User-Agent": "NovaJarvis/2.0"})
            with urllib.request.urlopen(req, timeout=120) as r:
                with open(ruta, "wb") as f:
                    f.write(r.read())
            # Mostrar en el visor de Jarvis (NO abrir ventana externa)
            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    "IMAGEN GENERADA", C["orb_speak"]))
            self.jarvis.root.after(50,
                lambda: self.jarvis._visor_add_imagen(ruta, prompt[:80]))
            # Registrar en el brain
            if hasattr(self.jarvis, "brain"):
                self.jarvis.brain.evento("imagen_generada", "jarvis",
                    {"prompt": prompt, "ruta": ruta})
                self.jarvis.brain.stat("jarvis", "imagenes", incrementar=1)
            return f"Imagen creada y mostrada en el visor: {prompt[:60]}", True
        except Exception as e:
            return f"Error generando imagen: {e}", False

    def crear_video_mini(self, tema, n_escenas="3"):
        """Genera N imagenes con narracion y las muestra como storyboard EN Jarvis.
        NO crea archivo de video, solo muestra las escenas en el visor."""
        try:
            n = max(1, min(6, int(n_escenas)))
        except:
            n = 3
        try:
            # Pedir a la IA que divida el tema en n escenas
            system = ("Eres un guionista. Devuelves EXACTAMENTE este formato "
                       "(sin marcadores extra):\n"
                       "ESCENA 1: [prompt visual en INGLES muy detallado] | "
                       "[narracion corta en espanol]\n"
                       "ESCENA 2: ... | ...\n")
            user = (f"Divide este tema en {n} escenas para un mini-video visual:\n"
                    f'"{tema}"\n\n'
                    f"Devuelve EXACTAMENTE {n} lineas 'ESCENA N: prompt | narracion'")
            resp, err = llamar_llm(self.jarvis.proveedor, self.jarvis.modelo,
                                    self.jarvis.api_key, system, user)
            if err:
                return f"Error generando guion: {err}", False

            # Parsear escenas
            escenas = []
            for linea in resp.split("\n"):
                m = re.match(r"ESCENA\s+\d+\s*:\s*(.+?)\s*\|\s*(.+)",
                              linea.strip(), re.IGNORECASE)
                if m:
                    escenas.append({
                        "prompt": m.group(1).strip(),
                        "narracion": m.group(2).strip(),
                    })
            if not escenas:
                return "La IA no devolvio escenas validas", False

            # Mostrar titulo en visor
            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    f"MINI-VIDEO: {tema[:40]}", C["orb_speak"]))
            self.jarvis.root.after(50,
                lambda: self.jarvis._visor_add_texto(
                    f"{len(escenas)} escenas | generando...", C["fg2"]))

            # Generar imagenes en background y mostrarlas
            def gen_todas():
                for i, esc in enumerate(escenas):
                    try:
                        slug = urllib.parse.quote(esc["prompt"] +
                            ", cinematic, dramatic lighting, 8k")
                        url = (f"https://image.pollinations.ai/prompt/{slug}"
                               f"?width=768&height=768&nologo=true"
                               f"&seed={int(time.time())%10000 + i}")
                        ruta = os.path.join(CAPTURAS,
                            f"video_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{i+1}.png")
                        req = urllib.request.Request(url,
                            headers={"User-Agent": "NovaJarvis/2.0"})
                        with urllib.request.urlopen(req, timeout=120) as r:
                            with open(ruta, "wb") as f:
                                f.write(r.read())
                        # Anadir al visor
                        self.jarvis.root.after(0,
                            lambda r=ruta, e=esc, n=i+1:
                                self.jarvis._visor_add_imagen(r,
                                    f"Escena {n}: {e['narracion'][:100]}"))
                        time.sleep(1)
                    except Exception as e:
                        print(f"Escena {i+1} falllo: {e}")

                # Al final, boton para "abrir en Studio"
                self.jarvis.root.after(0,
                    lambda: self.jarvis._visor_add_boton(
                        "Abrir NOVA Studio para editarlo",
                        lambda: self.abrir_modulo_nova("studio"),
                        C["orb_speak"]))

            threading.Thread(target=gen_todas, daemon=True).start()

            if hasattr(self.jarvis, "brain"):
                self.jarvis.brain.evento("mini_video_creado", "jarvis",
                    {"tema": tema, "n_escenas": len(escenas)})

            return (f"Mini-video de {len(escenas)} escenas generandose "
                    f"en el visor. Espera unos segundos por cada imagen."), True
        except Exception as e:
            return f"Error: {e}", False

    # === HERRAMIENTAS DE OFFICE ===
    def crear_word(self, tema, tipo="informe"):
        """Genera un documento Word con IA (delega a NOVA Office silenciosamente)."""
        try:
            # Reusar la logica de Office directamente
            docs_dir = os.path.join(PARENT, "Nova Office", "documentos")
            os.makedirs(docs_dir, exist_ok=True)

            system = ("Eres experto en redaccion. Devuelves documentos en "
                       "MARKDOWN. Estructura: # Titulo, ## Secciones, "
                       "**negrita**, - listas.")
            user = f"Crea un {tipo} sobre: {tema}. Devuelve solo el markdown."

            resp, err = llamar_llm(self.jarvis.proveedor, self.jarvis.modelo,
                                    self.jarvis.api_key, system, user)
            if err:
                return f"Error: {err}", False

            # Convertir a Word (importar aqui para no fallar si no esta)
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                doc = Document()
                for linea in resp.split("\n"):
                    s = linea.rstrip()
                    if not s:
                        doc.add_paragraph()
                    elif s.startswith("# "):
                        p = doc.add_heading(s[2:].strip(), level=0)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif s.startswith("## "):
                        doc.add_heading(s[3:].strip(), level=1)
                    elif s.startswith("### "):
                        doc.add_heading(s[4:].strip(), level=2)
                    elif s.lstrip().startswith(("- ", "* ")):
                        p = doc.add_paragraph(s.lstrip()[2:], style="List Bullet")
                    else:
                        doc.add_paragraph(s)

                slug = re.sub(r"[^\w\s-]", "", tema).strip().lower()
                slug = re.sub(r"[-\s]+", "_", slug)[:40] or "documento"
                nombre = slug + "_" + datetime.now().strftime("%Y%m%d_%H%M")
                ruta = os.path.join(docs_dir, f"{nombre}.docx")
                doc.save(ruta)

                # Mostrar preview del contenido en el visor
                self.jarvis.root.after(0,
                    lambda: self.jarvis._visor_add_titulo(
                        f"WORD: {tema[:40]}", C["accent2"]))
                self.jarvis.root.after(30,
                    lambda: self.jarvis._visor_add_texto(
                        resp[:600] + "..." if len(resp) > 600 else resp,
                        C["fg"], mono=False))
                self.jarvis.root.after(60,
                    lambda: self.jarvis._visor_add_boton(
                        f"Abrir {os.path.basename(ruta)}",
                        lambda r=ruta: os.startfile(r) if hasattr(os, "startfile") else None,
                        C["accent2"]))

                if hasattr(self.jarvis, "brain"):
                    self.jarvis.brain.evento("documento_creado", "jarvis",
                        {"formato": "word", "tema": tema, "ruta": ruta})

                return f"Documento Word creado: {ruta}", True
            except ImportError:
                return ("python-docx no instalado. Ejecuta: "
                        "pip install python-docx"), False
        except Exception as e:
            return f"Error creando Word: {e}", False

    def crear_excel(self, tema):
        """Genera un Excel con IA sobre un tema."""
        try:
            docs_dir = os.path.join(PARENT, "Nova Office", "documentos")
            os.makedirs(docs_dir, exist_ok=True)

            system = ("Devuelves datos CSV con separador '|'. Primera linea = "
                       "encabezados. Sin explicaciones, solo el CSV.")
            user = (f"Genera una tabla realista con datos sobre: {tema}. "
                    f"Usa '|' como separador. Minimo 8 filas de datos.")

            resp, err = llamar_llm(self.jarvis.proveedor, self.jarvis.modelo,
                                    self.jarvis.api_key, system, user)
            if err:
                return f"Error: {err}", False

            resp = re.sub(r"```\w*\n?", "", resp).strip("`\n ")
            filas = [l for l in resp.split("\n") if "|" in l]
            if not filas:
                return "IA no devolvio datos validos", False

            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, PatternFill, Alignment
                wb = Workbook()
                ws = wb.active
                data = [[c.strip() for c in f.split("|")] for f in filas]

                header_fill = PatternFill(start_color="60A5FA",
                                           end_color="60A5FA", fill_type="solid")
                header_font = Font(bold=True, color="FFFFFF")

                for r_idx, fila in enumerate(data, 1):
                    for c_idx, valor in enumerate(fila, 1):
                        cell = ws.cell(row=r_idx, column=c_idx, value=valor)
                        if r_idx == 1:
                            cell.fill = header_fill
                            cell.font = header_font
                            cell.alignment = Alignment(horizontal="center")

                slug = re.sub(r"[^\w\s-]", "", tema).strip().lower()
                slug = re.sub(r"[-\s]+", "_", slug)[:40] or "hoja"
                ruta = os.path.join(docs_dir,
                    f"{slug}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx")
                wb.save(ruta)

                # Preview: mostrar primera fila y filas en visor
                self.jarvis.root.after(0,
                    lambda: self.jarvis._visor_add_titulo(
                        f"EXCEL: {tema[:40]}", C["ok"]))
                preview = "\n".join(filas[:8])
                self.jarvis.root.after(30,
                    lambda: self.jarvis._visor_add_texto(preview,
                        C["fg"], mono=True))
                self.jarvis.root.after(60,
                    lambda: self.jarvis._visor_add_boton(
                        f"Abrir {os.path.basename(ruta)}",
                        lambda r=ruta: os.startfile(r) if hasattr(os, "startfile") else None,
                        C["ok"]))

                if hasattr(self.jarvis, "brain"):
                    self.jarvis.brain.evento("documento_creado", "jarvis",
                        {"formato": "excel", "tema": tema, "ruta": ruta})

                return f"Excel creado: {ruta} ({len(filas)} filas)", True
            except ImportError:
                return "openpyxl no instalado. pip install openpyxl", False
        except Exception as e:
            return f"Error creando Excel: {e}", False

    # === HERRAMIENTAS DE COACH ===
    def add_habito_coach(self, nombre, emoji="⭐"):
        """Agrega un habito a NOVA Coach."""
        try:
            coach_data = os.path.join(PARENT, "Nova Coach", "datos.json")
            if os.path.exists(coach_data):
                with open(coach_data, "r", encoding="utf-8") as f:
                    d = json.load(f)
            else:
                d = {"habitos": [], "registro": {}, "metas": [],
                     "perfil": {}, "planes": {}, "diario": []}

            d.setdefault("habitos", []).append({
                "emoji": emoji or "⭐",
                "nombre": nombre,
            })

            os.makedirs(os.path.dirname(coach_data), exist_ok=True)
            with open(coach_data, "w", encoding="utf-8") as f:
                json.dump(d, f, indent=2, ensure_ascii=False)

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo("HABITO ANADIDO", C["ok"]))
            self.jarvis.root.after(30,
                lambda: self.jarvis._visor_add_texto(
                    f"{emoji} {nombre}\n\nTotal habitos: {len(d['habitos'])}",
                    C["ok"]))

            if hasattr(self.jarvis, "brain"):
                self.jarvis.brain.evento("habito_creado", "jarvis",
                    {"nombre": nombre})

            return f"Habito anadido a Coach: {emoji} {nombre}", True
        except Exception as e:
            return f"Error: {e}", False

    def crear_meta_coach(self, texto):
        """Agrega una meta a NOVA Coach."""
        try:
            coach_data = os.path.join(PARENT, "Nova Coach", "datos.json")
            if os.path.exists(coach_data):
                with open(coach_data, "r", encoding="utf-8") as f:
                    d = json.load(f)
            else:
                d = {"habitos": [], "registro": {}, "metas": [],
                     "perfil": {}, "planes": {}, "diario": []}

            d.setdefault("metas", []).append({
                "texto": texto,
                "fecha_creacion": datetime.now().isoformat(),
                "completada": False,
            })

            os.makedirs(os.path.dirname(coach_data), exist_ok=True)
            with open(coach_data, "w", encoding="utf-8") as f:
                json.dump(d, f, indent=2, ensure_ascii=False)

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo("META CREADA", C["ok"]))
            self.jarvis.root.after(30,
                lambda: self.jarvis._visor_add_texto(
                    f"🎯 {texto}\n\nTotal metas activas: "
                    f"{sum(1 for m in d['metas'] if not m.get('completada'))}",
                    C["ok"]))

            return f"Meta creada: {texto}", True
        except Exception as e:
            return f"Error: {e}", False

    # === LECTURA DE PDFs (Reader silencioso) ===
    def resumir_pdf(self, ruta):
        """Lee un PDF y devuelve un resumen (mostrado en el visor)."""
        try:
            ruta = os.path.expanduser(ruta)
            if not os.path.exists(ruta):
                return f"No existe: {ruta}", False
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                return "PyPDF2 no instalado. pip install PyPDF2", False

            pdf = PdfReader(ruta)
            texto = ""
            for p in pdf.pages[:20]:  # max 20 paginas
                texto += p.extract_text() + "\n"

            if len(texto) < 50:
                return "PDF vacio o escaneado (sin OCR)", False

            texto = texto[:30000]  # limitar por si es enorme

            system = ("Resume documentos en espanol de forma clara y "
                       "estructurada con puntos clave.")
            user = f"Resume este texto en 5-7 puntos clave:\n\n{texto}"

            resp, err = llamar_llm(self.jarvis.proveedor, self.jarvis.modelo,
                                    self.jarvis.api_key, system, user)
            if err:
                return f"Error resumiendo: {err}", False

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    f"RESUMEN PDF: {os.path.basename(ruta)[:40]}",
                    C["accent2"]))
            self.jarvis.root.after(30,
                lambda: self.jarvis._visor_add_texto(resp, C["fg"]))

            if hasattr(self.jarvis, "brain"):
                self.jarvis.brain.evento("pdf_resumido", "jarvis",
                    {"archivo": os.path.basename(ruta)})

            return f"PDF resumido en el visor: {os.path.basename(ruta)}", True
        except Exception as e:
            return f"Error: {e}", False

    # === INFO ===
    def decir_hora(self):
        ahora = datetime.now()
        return ahora.strftime("Son las %H:%M del %A %d de %B de %Y"), True

    def decir_clima(self, ciudad="Madrid"):
        try:
            # wttr.in devuelve clima sin API key
            url = f"https://wttr.in/{urllib.parse.quote(ciudad)}?format=%C+%t+%h+%w&lang=es"
            req = urllib.request.Request(url,
                headers={"User-Agent": "curl/7.79.1"})
            with urllib.request.urlopen(req, timeout=10) as r:
                info = r.read().decode("utf-8").strip()
            return f"Clima en {ciudad}: {info}", True
        except Exception as e:
            return f"No pude obtener el clima: {e}", False

    def calcular_expresion(self, expr):
        try:
            # Solo permitir caracteres seguros
            if not re.match(r"^[\d\s\+\-\*\/\.\(\)\%\*\*]+$", expr):
                return "Expresion no permitida (solo numeros y operadores)", False
            resultado = eval(expr, {"__builtins__": {}}, {})
            return f"{expr} = {resultado}", True
        except Exception as e:
            return f"Error de calculo: {e}", False

    # === MEMORIA ===
    def recordar_hecho(self, hecho):
        # Guarda en local (Jarvis)
        self.jarvis.memoria.setdefault("hechos", []).append({
            "texto": hecho,
            "fecha": datetime.now().isoformat(),
        })
        guardar_json(MEMORIA, self.jarvis.memoria)
        # Y tambien en el cerebro compartido (todos los modulos NOVA)
        self.jarvis.brain.recordar(hecho, modulo="jarvis")
        return f"He recordado: {hecho}", True

    def olvidar_hecho(self, indice):
        try:
            idx = int(indice)
            hechos = self.jarvis.memoria.get("hechos", [])
            if 0 <= idx < len(hechos):
                borrado = hechos.pop(idx)
                guardar_json(MEMORIA, self.jarvis.memoria)
                return f"Olvidado: {borrado['texto']}", True
            return "Indice no valido", False
        except Exception as e:
            return f"Error: {e}", False

    def listar_hechos(self):
        hechos = self.jarvis.memoria.get("hechos", [])
        if not hechos:
            return "No recuerdo ningun hecho aun.", True
        lineas = [f"[{i}] {h['texto']}" for i, h in enumerate(hechos)]
        return "Recuerdo estos hechos sobre ti:\n" + "\n".join(lineas), True

    # === RECORDATORIOS ===
    def crear_recordatorio(self, texto, minutos):
        try:
            minutos = int(minutos)
            cuando = datetime.now() + timedelta(minutes=minutos)
            self.jarvis.recordatorios.append({
                "texto": texto,
                "hora": cuando.isoformat(),
                "notificado": False,
            })
            guardar_json(RECORDATORIOS, self.jarvis.recordatorios)
            return f"Recordatorio creado para dentro de {minutos} min: {texto}", True
        except Exception as e:
            return f"Error: {e}", False

    def listar_recordatorios(self):
        activos = [r for r in self.jarvis.recordatorios if not r.get("notificado")]
        if not activos:
            return "No tienes recordatorios activos.", True
        lineas = []
        for r in activos:
            try:
                cuando = datetime.fromisoformat(r["hora"])
                delta = cuando - datetime.now()
                mins = int(delta.total_seconds() / 60)
                lineas.append(f"  - En {mins} min: {r['texto']}")
            except Exception:
                lineas.append(f"  - {r['texto']}")
        return "Tienes estos recordatorios:\n" + "\n".join(lineas), True

    # === MODULOS NOVA ===
    # ============================================================
    # SUPERPODERES: VISION
    # ============================================================
    def analizar_pantalla(self, pregunta="Que hay en la pantalla?"):
        """Captura la pantalla y la analiza con IA multimodal."""
        if not PIL_OK:
            return "Necesitas Pillow: pip install Pillow", False
        try:
            img = ImageGrab.grab()
            ruta = os.path.join(CAPTURAS,
                f"vision_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
            img.save(ruta)

            # Mostrar la captura en el visor
            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    "ANALIZANDO PANTALLA", C["orb_speak"]))
            self.jarvis.root.after(30,
                lambda: self.jarvis._visor_add_imagen(ruta,
                    "Captura en analisis..."))

            # Enviar a la IA con vision
            resp, err = llamar_llm(self.jarvis.proveedor, self.jarvis.modelo,
                self.jarvis.api_key,
                "Eres un experto en analizar imagenes de pantalla. "
                "Describes con detalle lo que ves y respondes en espanol.",
                pregunta, imagen=ruta, timeout=90)
            if err:
                return f"Error analizando: {err}", False

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_texto(resp, C["fg"]))

            if hasattr(self.jarvis, "brain"):
                self.jarvis.brain.evento("pantalla_analizada", "jarvis",
                    {"pregunta": pregunta})
            return f"Pantalla analizada. Ver resultado en visor.", True
        except Exception as e:
            return f"Error: {e}", False

    def analizar_imagen(self, ruta, pregunta="Describe esta imagen"):
        """Analiza una imagen (ruta local) con IA multimodal."""
        try:
            ruta = os.path.expanduser(ruta)
            if not os.path.exists(ruta):
                return f"No existe: {ruta}", False

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    f"ANALIZANDO: {os.path.basename(ruta)}", C["orb_speak"]))
            self.jarvis.root.after(30,
                lambda: self.jarvis._visor_add_imagen(ruta, pregunta))

            resp, err = llamar_llm(self.jarvis.proveedor, self.jarvis.modelo,
                self.jarvis.api_key,
                "Eres un experto en analizar imagenes. Respondes en espanol "
                "con detalle y precision.",
                pregunta, imagen=ruta, timeout=90)
            if err:
                return f"Error: {err}", False

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_texto(resp, C["fg"]))
            return "Imagen analizada. Ver resultado en visor.", True
        except Exception as e:
            return f"Error: {e}", False

    def extraer_texto_imagen(self, ruta):
        """Extrae texto de una imagen usando la IA (OCR con IA)."""
        return self.analizar_imagen(ruta,
            "Extrae y transcribe TODO el texto visible en la imagen. "
            "Devuelve solo el texto tal cual, sin comentarios.")

    # ============================================================
    # SUPERPODERES: AGENTE (planificacion multi-paso)
    # ============================================================
    def modo_agente(self, objetivo):
        """
        Modo Agente: la IA planifica varios pasos y los ejecuta uno a uno.
        Ejemplo objetivo: "investiga IA en 2026 y crea un word con el resumen"
        """
        self.jarvis.root.after(0,
            lambda: self.jarvis._visor_add_titulo(
                "MODO AGENTE ACTIVADO", C["accent2"]))
        self.jarvis.root.after(30,
            lambda: self.jarvis._visor_add_texto(
                f"Objetivo: {objetivo}\n\nPlanificando pasos...", C["fg"]))

        # Pedir a la IA que descomponga el objetivo en pasos
        system = ("Eres un planificador. Descompones objetivos complejos en "
                   "3-5 pasos ejecutables. Cada paso debe ser UNA accion "
                   "concreta que Jarvis puede hacer con sus herramientas.\n\n"
                   "Herramientas disponibles: buscar_wikipedia, buscar_web, "
                   "generar_imagen, crear_video_mini, crear_word, crear_excel, "
                   "add_habito_coach, crear_meta_coach, resumir_pdf, "
                   "recordar_hecho, calcular_expresion, decir_clima, "
                   "abrir_modulo_nova.\n\n"
                   "Devuelve EXACTAMENTE este formato:\n"
                   "PASO 1: [descripcion] | [[TOOL:nombre|arg1|arg2]]\n"
                   "PASO 2: [descripcion] | [[TOOL:nombre|arg1|arg2]]\n"
                   "...")
        user = f"Objetivo: {objetivo}\n\nDescompon en pasos ejecutables."

        try:
            plan, err = llamar_llm(self.jarvis.proveedor, self.jarvis.modelo,
                                    self.jarvis.api_key, system, user)
            if err:
                return f"Error planificando: {err}", False

            # Parsear pasos
            pasos = []
            for linea in plan.split("\n"):
                m = re.match(r"PASO\s+\d+\s*:\s*(.+?)\s*\|\s*(.+)",
                              linea.strip(), re.IGNORECASE)
                if m:
                    desc = m.group(1).strip()
                    tool_str = m.group(2).strip()
                    tool_m = TOOL_RE.search(tool_str)
                    if tool_m:
                        nombre = tool_m.group(1)
                        args_raw = tool_m.group(2)
                        args = [a for a in args_raw.split("|") if a] if args_raw else []
                        pasos.append({"desc": desc, "tool": nombre, "args": args})
            if not pasos:
                return "El agente no genero pasos ejecutables validos", False

            # Mostrar el plan en el visor
            plan_txt = f"PLAN ({len(pasos)} pasos):\n\n"
            for i, p in enumerate(pasos, 1):
                plan_txt += f"{i}. {p['desc']}\n   -> {p['tool']}({','.join(p['args'])[:60]})\n\n"
            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_texto(plan_txt,
                    C["accent2"], mono=True))

            # Ejecutar pasos en background
            def ejecutar():
                for i, p in enumerate(pasos, 1):
                    self.jarvis.root.after(0,
                        lambda i=i, p=p: self.jarvis._visor_add_titulo(
                            f"PASO {i}/{len(pasos)}: {p['tool']}",
                            C["orb_tool"]))
                    metodo = getattr(self, p["tool"], None)
                    if not metodo:
                        self.jarvis.root.after(0,
                            lambda p=p: self.jarvis._visor_add_texto(
                                f"Herramienta desconocida: {p['tool']}",
                                C["warn"]))
                        continue
                    try:
                        resultado, ok = metodo(*p["args"]) if p["args"] else metodo()
                        color = C["ok"] if ok else C["warn"]
                        # Truncar resultado si es muy largo
                        r_show = resultado[:300] + "..." if len(resultado) > 300 else resultado
                        self.jarvis.root.after(0,
                            lambda r=r_show, c=color: self.jarvis._visor_add_texto(
                                r, c))
                    except Exception as e:
                        self.jarvis.root.after(0,
                            lambda e=e: self.jarvis._visor_add_texto(
                                f"Error: {e}", C["err"]))
                    time.sleep(0.5)  # pausa entre pasos

                self.jarvis.root.after(0,
                    lambda: self.jarvis._visor_add_titulo(
                        "AGENTE COMPLETO", C["ok"]))
                if hasattr(self.jarvis, "brain"):
                    self.jarvis.brain.evento("agente_ejecutado", "jarvis",
                        {"objetivo": objetivo, "pasos": len(pasos)})

            threading.Thread(target=ejecutar, daemon=True).start()
            return f"Agente ejecutando {len(pasos)} pasos. Ver progreso en visor.", True
        except Exception as e:
            return f"Error del agente: {e}", False

    # ============================================================
    # SUPERPODERES: AUTOMATIZACIONES / RUTINAS
    # ============================================================
    def crear_rutina(self, nombre, hora_hhmm, acciones):
        """
        Crea una rutina programada.
        hora_hhmm: "09:00" o "cada 30min"
        acciones: texto con lo que debe hacer (ej: "dame el clima y mis metas")
        """
        try:
            self.jarvis.rutinas.append({
                "nombre": nombre,
                "hora": hora_hhmm,
                "acciones": acciones,
                "ultima_ejecucion": None,
                "activa": True,
                "creada": datetime.now().isoformat(),
            })
            guardar_json(RUTINAS, self.jarvis.rutinas)

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    "RUTINA CREADA", C["ok"]))
            self.jarvis.root.after(30,
                lambda: self.jarvis._visor_add_texto(
                    f"{nombre}\nHora: {hora_hhmm}\nAcciones: {acciones}",
                    C["ok"]))

            return f"Rutina '{nombre}' programada para {hora_hhmm}", True
        except Exception as e:
            return f"Error: {e}", False

    def listar_rutinas(self):
        rutinas = [r for r in self.jarvis.rutinas if r.get("activa", True)]
        if not rutinas:
            return "No tienes rutinas activas.", True
        lineas = ["Rutinas activas:"]
        for r in rutinas:
            lineas.append(f"  - {r['nombre']} @ {r['hora']}: {r['acciones'][:60]}")
        return "\n".join(lineas), True

    def borrar_rutina(self, nombre):
        antes = len(self.jarvis.rutinas)
        self.jarvis.rutinas = [r for r in self.jarvis.rutinas
                                 if r["nombre"].lower() != nombre.lower()]
        guardar_json(RUTINAS, self.jarvis.rutinas)
        borradas = antes - len(self.jarvis.rutinas)
        return f"Borradas {borradas} rutinas con nombre '{nombre}'", True

    # ============================================================
    # SUPERPODERES: RAG (chat con tus documentos)
    # ============================================================
    def rag_indexar(self, ruta_carpeta):
        """Indexa una carpeta con documentos para chatear con ellos."""
        if not RAG_OK:
            return ("Instala PyPDF2 para RAG: pip install PyPDF2 python-docx"), False
        try:
            if not hasattr(self.jarvis, "rag") or self.jarvis.rag is None:
                self.jarvis.rag = NovaRAG()

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    "INDEXANDO DOCUMENTOS", C["accent2"]))

            def prog(actual, total, nombre):
                self.jarvis.root.after(0,
                    lambda a=actual, t=total, n=nombre:
                        self.jarvis._visor_add_texto(
                            f"[{a}/{t}] {n}", C["fg2"]))

            docs, chunks = self.jarvis.rag.indexar(ruta_carpeta,
                reemplazar=False, progress=prog)

            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_texto(
                    f"OK - Indexados {docs} documentos ({chunks} fragmentos)",
                    C["ok"]))

            if hasattr(self.jarvis, "brain"):
                self.jarvis.brain.evento("rag_indexado", "jarvis",
                    {"carpeta": ruta_carpeta, "docs": docs})

            return f"Indexados {docs} documentos, {chunks} fragmentos.", True
        except Exception as e:
            return f"Error indexando: {e}", False

    def rag_preguntar(self, consulta):
        """Pregunta usando tus documentos como fuente (RAG)."""
        if not RAG_OK:
            return "Instala RAG: pip install PyPDF2 python-docx", False
        try:
            if not hasattr(self.jarvis, "rag") or self.jarvis.rag is None:
                self.jarvis.rag = NovaRAG()

            if not self.jarvis.rag.chunks:
                return ("No hay documentos indexados. Usa rag_indexar primero "
                        "con una ruta de carpeta."), False

            resp, fragmentos = self.jarvis.rag.preguntar(
                consulta, llamar_llm,
                self.jarvis.proveedor, self.jarvis.modelo,
                self.jarvis.api_key)

            # Mostrar en el visor
            self.jarvis.root.after(0,
                lambda: self.jarvis._visor_add_titulo(
                    "RAG: Respuesta con fuentes", C["accent2"]))
            self.jarvis.root.after(30,
                lambda: self.jarvis._visor_add_texto(resp, C["fg"]))

            if fragmentos:
                self.jarvis.root.after(50,
                    lambda: self.jarvis._visor_add_titulo(
                        f"FUENTES CONSULTADAS ({len(fragmentos)})", C["fg2"]))
                for f in fragmentos:
                    self.jarvis.root.after(80,
                        lambda f=f: self.jarvis._visor_add_texto(
                            f"[{f['score']}] {f['fuente']}\n{f['texto'][:200]}...",
                            C["fg2"], mono=True))

            return resp, True
        except Exception as e:
            return f"Error: {e}", False

    def rag_resumen(self):
        """Muestra que hay indexado actualmente."""
        if not RAG_OK:
            return "RAG no disponible", False
        if not hasattr(self.jarvis, "rag") or self.jarvis.rag is None:
            self.jarvis.rag = NovaRAG()
        r = self.jarvis.rag.resumen()
        return (f"Indice RAG: {r['docs']} documentos, {r['chunks']} fragmentos. "
                f"Embeddings: {'SI' if r['usa_embeddings'] else 'NO (usa TF-IDF)'}."), True

    def rag_limpiar(self):
        """Borra el indice RAG."""
        if hasattr(self.jarvis, "rag") and self.jarvis.rag:
            self.jarvis.rag.limpiar()
            return "Indice RAG borrado", True
        return "No hay indice", True

    # ============================================================
    # SUPERPODERES: 15 HERRAMIENTAS NUEVAS
    # ============================================================

    # === CONTROL DE MUSICA ===
    def musica_play_pause(self):
        """Pausa o reanuda musica."""
        if not MUSICA_OK:
            return "Musica no disponible", False
        ok = nova_musica.play_pause()
        return "Play/Pause enviado" if ok else "No se pudo enviar", ok

    def musica_siguiente(self):
        """Siguiente cancion."""
        if not MUSICA_OK: return "Musica no disponible", False
        nova_musica.siguiente()
        return "Siguiente cancion", True

    def musica_anterior(self):
        """Cancion anterior."""
        if not MUSICA_OK: return "Musica no disponible", False
        nova_musica.anterior()
        return "Cancion anterior", True

    def musica_volumen(self, direccion="subir", pasos="3"):
        """Sube o baja el volumen. direccion: 'subir'|'bajar'|'mute'."""
        if not MUSICA_OK: return "Musica no disponible", False
        try:
            n = int(pasos)
        except:
            n = 3
        d = direccion.lower().strip()
        if d in ("subir", "up", "arriba", "mas"):
            nova_musica.volumen_subir(n)
            return f"Volumen +{n}", True
        if d in ("bajar", "down", "abajo", "menos"):
            nova_musica.volumen_bajar(n)
            return f"Volumen -{n}", True
        if d in ("mute", "silencio", "silenciar"):
            nova_musica.silenciar()
            return "Silenciado (toggle)", True
        return f"Direccion invalida: {direccion}", False

    def musica_reproducir(self, busqueda):
        """Busca y reproduce en YouTube Music."""
        if not MUSICA_OK: return "Musica no disponible", False
        r = nova_musica.buscar_youtube_music(busqueda)
        return r, True

    def musica_mood(self, mood):
        """Reproduce musica segun un mood (relax, concentracion, energia, fiesta...)."""
        if not MUSICA_OK: return "Musica no disponible", False
        r = nova_musica.reproducir_mood(mood)
        return r, True

    def musica_spotify(self, busqueda=""):
        """Abre Spotify (opcionalmente buscando algo)."""
        if not MUSICA_OK: return "Musica no disponible", False
        if busqueda:
            return nova_musica.buscar_spotify_web(busqueda), True
        return nova_musica.abrir_spotify_app(), True

    # === COMUNICACION ===
    def abrir_whatsapp_web(self, mensaje_a=""):
        """Abre WhatsApp Web. Si mensaje_a='numero+mensaje' abre chat con
           un numero. Ej: '34600000000+Hola que tal'."""
        import webbrowser
        if "+" in mensaje_a:
            partes = mensaje_a.split("+", 1)
            numero = re.sub(r"\D", "", partes[0])
            msg = urllib.parse.quote(partes[1] if len(partes) > 1 else "")
            url = f"https://wa.me/{numero}?text={msg}"
        else:
            url = "https://web.whatsapp.com/"
        webbrowser.open(url)
        return f"Abriendo WhatsApp: {url}", True

    def abrir_correo(self, destinatario="", asunto="", cuerpo=""):
        """Abre tu app de correo predeterminada componiendo un email."""
        import webbrowser
        params = []
        if asunto:
            params.append(f"subject={urllib.parse.quote(asunto)}")
        if cuerpo:
            params.append(f"body={urllib.parse.quote(cuerpo)}")
        url = f"mailto:{destinatario}"
        if params:
            url += "?" + "&".join(params)
        webbrowser.open(url)
        return f"Correo compuesto para: {destinatario or '(sin destino)'}", True

    def abrir_gmail_web(self, destinatario="", asunto="", cuerpo=""):
        """Compone en Gmail Web (sin abrir app local)."""
        import webbrowser
        url = ("https://mail.google.com/mail/?view=cm&fs=1"
               f"&to={urllib.parse.quote(destinatario)}"
               f"&su={urllib.parse.quote(asunto)}"
               f"&body={urllib.parse.quote(cuerpo)}")
        webbrowser.open(url)
        return f"Gmail compuesto para: {destinatario}", True

    # === NOTAS RAPIDAS ===
    def nota_rapida(self, texto):
        """Guarda una nota rapida en notas.txt."""
        ruta = os.path.join(APP_DIR, "notas_rapidas.txt")
        try:
            with open(ruta, "a", encoding="utf-8") as f:
                f.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] {texto}\n")
            return f"Nota guardada: {texto[:60]}", True
        except Exception as e:
            return f"Error: {e}", False

    def leer_notas(self, cantidad="10"):
        """Lee las ultimas N notas rapidas."""
        ruta = os.path.join(APP_DIR, "notas_rapidas.txt")
        if not os.path.exists(ruta):
            return "Aun no tienes notas.", True
        try:
            n = int(cantidad)
        except:
            n = 10
        try:
            with open(ruta, "r", encoding="utf-8") as f:
                lineas = f.readlines()
            ultimas = lineas[-n:]
            return "Tus ultimas notas:\n" + "".join(ultimas), True
        except Exception as e:
            return f"Error: {e}", False

    # === CALENDARIO / GOOGLE APPS ===
    def abrir_calendar(self, cuando=""):
        """Abre Google Calendar (opcionalmente crea un evento nuevo)."""
        import webbrowser
        if cuando:
            url = ("https://calendar.google.com/calendar/render?action=TEMPLATE"
                   f"&text={urllib.parse.quote(cuando)}")
        else:
            url = "https://calendar.google.com/"
        webbrowser.open(url)
        return f"Google Calendar abierto", True

    def buscar_maps(self, lugar):
        """Busca un lugar en Google Maps."""
        import webbrowser
        url = f"https://www.google.com/maps/search/{urllib.parse.quote(lugar)}"
        webbrowser.open(url)
        return f"Maps: {lugar}", True

    # === GITHUB / DEV ===
    def buscar_github(self, query):
        """Busca repos en GitHub."""
        import webbrowser
        url = f"https://github.com/search?q={urllib.parse.quote(query)}&type=repositories"
        webbrowser.open(url)
        return f"GitHub: {query}", True

    def abrir_stackoverflow(self, pregunta):
        """Busca en Stack Overflow."""
        import webbrowser
        url = f"https://stackoverflow.com/search?q={urllib.parse.quote(pregunta)}"
        webbrowser.open(url)
        return f"Stack Overflow: {pregunta}", True

    # === VISION CONTINUA ===
    def vision_continua_iniciar(self, intervalo_seg="10"):
        """Activa el modo vision continua: cada N seg guarda una captura
        que la IA puede consultar."""
        try:
            n = max(3, int(intervalo_seg))
        except:
            n = 10
        self.jarvis.vision_continua_intervalo = n
        self.jarvis.vision_continua_activa = True
        self.jarvis.root.after(500, self.jarvis._loop_vision_continua)
        return f"Vision continua ACTIVADA cada {n} segundos. Ahora sabre que ves.", True

    def vision_continua_parar(self):
        """Desactiva el modo vision continua."""
        self.jarvis.vision_continua_activa = False
        return "Vision continua desactivada.", True

    def que_veo_ahora(self, pregunta="Describe brevemente que hay en la pantalla"):
        """Captura la pantalla AHORA y la analiza."""
        return self.analizar_pantalla(pregunta)

    def abrir_modulo_nova(self, nombre):
        mapa = {
            "reader": ("Nova Reader", "nova_reader.py"),
            "studio": ("Nova Studio", "nova_studio.py"),
            "coach": ("Nova Coach", "nova_coach.py"),
            "office": ("Nova Office", "nova_office.py"),
            "search": ("Nova Search", "nova_search.py"),
            "game": ("Nova Game Master", "nova_game_master.py"),
            "gm": ("Nova Game Master", "nova_game_master.py"),
            "hub": ("Nova Brain Hub", "nova_brain_hub.py"),
            "mini": ("Nova Mini-Beta", "nova_mini.py"),
            "orb": ("Nova Orb", "nova_orb.py"),
        }
        clave = nombre.lower().strip()
        for k, (carpeta, archivo) in mapa.items():
            if k in clave:
                ruta = os.path.join(PARENT, carpeta, archivo)
                if os.path.exists(ruta):
                    subprocess.Popen([sys.executable, ruta],
                                      cwd=os.path.dirname(ruta))
                    return f"Abriendo modulo NOVA {k}", True
                return f"No encuentro {carpeta}/{archivo}", False
        return f"Modulo desconocido: {nombre}. Prueba: reader, studio, coach, office, search, game, hub, mini, orb", False


# ============================================================
# PARSER DE HERRAMIENTAS EN LA RESPUESTA DE LA IA
# ============================================================
# Formato esperado en la respuesta: [[TOOL:nombre|arg1|arg2]]
TOOL_RE = re.compile(r"\[\[TOOL:([a-zA-Z_]+)((?:\|[^\]]*)*?)\]\]")

def parsear_tools(texto):
    """Devuelve lista de (herramienta, [args]) que aparecen en texto."""
    encontrados = []
    for m in TOOL_RE.finditer(texto):
        nombre = m.group(1)
        args_raw = m.group(2)
        args = [a for a in args_raw.split("|") if a] if args_raw else []
        encontrados.append((nombre, args, m.group(0)))
    return encontrados

# ============================================================
# ESFERA ANIMADA (Iron Man style)
# ============================================================
class Esfera:
    def __init__(self, canvas, cx, cy, r=60):
        self.canvas = canvas
        self.cx, self.cy, self.r = cx, cy, r
        self.estado = "idle"
        self.fase = 0
        self._crear()

    def _crear(self):
        # Anillo exterior
        self.anillo = self.canvas.create_oval(
            self.cx - self.r - 10, self.cy - self.r - 10,
            self.cx + self.r + 10, self.cy + self.r + 10,
            outline=C["orb_idle"], width=2, fill="")
        # Esfera principal
        self.orb = self.canvas.create_oval(
            self.cx - self.r, self.cy - self.r,
            self.cx + self.r, self.cy + self.r,
            fill=C["orb_idle"], outline="")
        # Nucleo interior brillante
        self.nucleo = self.canvas.create_oval(
            self.cx - self.r // 2, self.cy - self.r // 2,
            self.cx + self.r // 2, self.cy + self.r // 2,
            fill="#ffffff", outline="")

    def set_estado(self, est):
        self.estado = est
        color = C.get(f"orb_{est}", C["orb_idle"])
        self.canvas.itemconfig(self.orb, fill=color)
        self.canvas.itemconfig(self.anillo, outline=color)

    def animar(self):
        self.fase += 1
        # Pulsacion segun estado
        if self.estado == "listen":
            # Rapida
            expansion = int(5 * abs((self.fase % 20) - 10) / 10)
        elif self.estado == "think":
            # Media
            expansion = int(8 * abs((self.fase % 40) - 20) / 20)
        elif self.estado == "speak":
            # Rapida y grande
            expansion = int(12 * abs((self.fase % 20) - 10) / 10)
        elif self.estado == "tool":
            expansion = int(10 * abs((self.fase % 30) - 15) / 15)
        else:
            expansion = int(3 * abs((self.fase % 60) - 30) / 30)
        r = self.r + expansion
        self.canvas.coords(self.orb,
            self.cx - r, self.cy - r,
            self.cx + r, self.cy + r)
        r2 = self.r + 10 + expansion
        self.canvas.coords(self.anillo,
            self.cx - r2, self.cy - r2,
            self.cx + r2, self.cy + r2)


# ============================================================
# TTS con edge-tts
# ============================================================
def hablar_voz(texto, voz_id):
    if not EDGE_OK:
        return False
    ruta = os.path.join(CAPTURAS, "_ultimo.mp3")
    try:
        async def gen():
            comm = edge_tts.Communicate(texto, voz_id)
            await comm.save(ruta)
        loop = asyncio.new_event_loop()
        loop.run_until_complete(gen())
        loop.close()
        if PYGAME_OK and _init_pygame():
            pygame.mixer.music.load(ruta)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                time.sleep(0.1)
        else:
            try: os.startfile(ruta)
            except AttributeError: pass
        return True
    except Exception as e:
        print(f"TTS error: {e}")
        return False


# ============================================================
# APLICACION PRINCIPAL
# ============================================================
class NovaJarvis:
    def __init__(self, root):
        self.root = root
        self.root.title("NOVA Jarvis - Asistente todopoderoso")
        self.root.geometry("1200x820")
        self.root.configure(bg=C["bg"])

        # Config
        self.cfg = cargar_json(CONFIG, {"api_keys": {}})
        self.proveedor = self.cfg.get("jarvis_proveedor", "Groq")
        self.modelo = self.cfg.get("jarvis_modelo", "llama-3.3-70b-versatile")
        self.api_key = self.cfg.get("api_keys", {}).get(self.proveedor, "")
        self.voz = self.cfg.get("jarvis_voz", "es-ES-ElviraNeural")
        self.voz_activada = self.cfg.get("jarvis_voz_activada", True)
        self.tools_habilitadas = self.cfg.get("jarvis_tools", True)

        # Memoria (dual: local + cerebro global compartido)
        self.memoria = cargar_json(MEMORIA,
            {"hechos": [], "conversaciones": []})
        self.brain = Brain()  # cerebro compartido con todo el ecosistema NOVA
        self.recordatorios = cargar_json(RECORDATORIOS, [])
        self.rutinas = cargar_json(RUTINAS, [])  # rutinas programadas
        # RAG y voz pro (cargados perezosamente)
        self.rag = None
        self.voz_pro = None
        self.voz_pro_activa = False
        # Personalidad activa
        self.personalidad = self.cfg.get("jarvis_personalidad", "formal")
        # Vision continua
        self.vision_continua_activa = False
        self.vision_continua_intervalo = 10
        self.vision_continua_ultima_captura = None
        self.vision_continua_ultimo_analisis = ""
        self.historial_actual = []
        self.procesando = False
        self.modo_continuo = False
        # Contador de conversaciones
        self.brain.stat("jarvis", "sesiones_iniciadas", incrementar=1)

        # Herramientas
        self.herramientas = Herramientas(self)

        self._build()
        self._mensaje_bienvenida()
        # Check recordatorios cada 30s
        self.root.after(2000, self._check_recordatorios)
        # Check rutinas cada minuto
        self.root.after(5000, self._check_rutinas)
        # Comprobar actualizaciones a los 3 seg del arranque (silencioso)
        self.root.after(3000, self._comprobar_update_silencioso)
        # Animacion esfera cada 50ms
        self.root.after(50, self._loop_animacion)

    def _build(self):
        # === HEADER ===
        header = tk.Frame(self.root, bg=C["bg2"], height=64)
        header.pack(fill="x"); header.pack_propagate(False)

        tk.Label(header, text="NOVA JARVIS",
            bg=C["bg2"], fg=C["jarvis"],
            font=("Segoe UI", 18, "bold")).pack(side="left", padx=20)
        tk.Label(header, text="Asistente todopoderoso con IA",
            bg=C["bg2"], fg=C["fg2"],
            font=("Segoe UI", 10, "italic")).pack(side="left", padx=4)

        for txt, cmd in [
            ("Config IA", self._config),
            ("Voz", self._config_voz),
            ("🎭 Personalidad", self._config_personalidad),
            ("Memoria", self._ver_memoria),
            ("Herramientas", self._ver_herramientas),
            ("🆕 Update", self._comprobar_update_manual),
            ("Ayuda", self._ver_ayuda),
        ]:
            tk.Button(header, text=txt, command=cmd,
                bg=C["bg3"], fg=C["fg"], relief="flat",
                font=("Segoe UI", 9, "bold"), cursor="hand2",
                padx=10, pady=4).pack(side="right", padx=3, pady=12)

        # === CUERPO ===
        body = tk.Frame(self.root, bg=C["bg"])
        body.pack(fill="both", expand=True)

        # --- Panel izquierdo: esfera + botones ---
        left = tk.Frame(body, bg=C["bg2"], width=280)
        left.pack(side="left", fill="y")
        left.pack_propagate(False)

        # Esfera
        self.canvas = tk.Canvas(left, bg=C["bg2"],
            highlightthickness=0, width=260, height=260)
        self.canvas.pack(pady=(30, 10))
        self.esfera = Esfera(self.canvas, 130, 130, r=60)

        self.lbl_estado = tk.Label(left, text="Listo",
            bg=C["bg2"], fg=C["fg"],
            font=("Segoe UI", 12, "bold"))
        self.lbl_estado.pack(pady=6)

        # Botones voz
        self.btn_voz = tk.Button(left, text="Voz: ON" if self.voz_activada else "Voz: OFF",
            command=self._toggle_voz,
            bg=C["ok"] if self.voz_activada else C["bg3"],
            fg="white", relief="flat",
            font=("Segoe UI", 10, "bold"), cursor="hand2",
            pady=8)
        self.btn_voz.pack(fill="x", padx=20, pady=4)

        self.btn_continuo = tk.Button(left, text="Modo escucha continua",
            command=self._toggle_continuo,
            bg=C["bg3"], fg=C["fg"], relief="flat",
            font=("Segoe UI", 10, "bold"), cursor="hand2",
            pady=8)
        self.btn_continuo.pack(fill="x", padx=20, pady=4)

        tk.Button(left, text="Escuchar una vez",
            command=self._escuchar_una,
            bg=C["bg3"], fg=C["fg"], relief="flat",
            font=("Segoe UI", 10), cursor="hand2",
            pady=8).pack(fill="x", padx=20, pady=4)

        # BOTON: Wake Word (voz profesional)
        self.btn_wake = tk.Button(left,
            text="Wake 'Hey Jarvis': OFF",
            command=self._toggle_wake_word,
            bg=C["bg3"], fg=C["fg"], relief="flat",
            font=("Segoe UI", 10, "bold"), cursor="hand2",
            pady=8)
        self.btn_wake.pack(fill="x", padx=20, pady=4)

        # Toggle tools
        self.btn_tools = tk.Button(left,
            text="Herramientas: " + ("ON" if self.tools_habilitadas else "OFF"),
            command=self._toggle_tools,
            bg=C["tool"] if self.tools_habilitadas else C["bg3"],
            fg="white", relief="flat",
            font=("Segoe UI", 10, "bold"), cursor="hand2",
            pady=8)
        self.btn_tools.pack(fill="x", padx=20, pady=4)

        # Quick actions
        tk.Label(left, text="ACCIONES RAPIDAS",
            bg=C["bg2"], fg=C["fg2"],
            font=("Segoe UI", 8, "bold")).pack(pady=(20, 4))
        for txt, cmd in [
            ("Que hora es?", lambda: self._enviar_directo("Que hora es?")),
            ("Clima Madrid", lambda: self._enviar_directo("Como esta el clima en Madrid?")),
            ("Ver pantalla", lambda: self._enviar_directo(
                "Analiza mi pantalla y dime que ves")),
            ("Generar imagen", lambda: self._enviar_directo(
                "Genera una imagen de un paisaje futurista")),
            ("Video 3 escenas", lambda: self._enviar_directo(
                "Crea un mini-video de 3 escenas sobre astronautas")),
            ("Word rapido", lambda: self._enviar_directo(
                "Crea un word con un ensayo corto sobre la IA")),
            ("Modo AGENTE", lambda: self._enviar_directo(
                "Usa modo agente para: investigar la historia de Espana en el siglo XX, "
                "crear un word con el resumen y una imagen ilustrativa")),
            ("🎵 Musica chill", lambda: self._enviar_directo(
                "Pon musica chill de fondo")),
            ("👁️ Que veo?", lambda: self._enviar_directo(
                "Que ves en mi pantalla ahora?")),
            ("📝 Nota rapida", lambda: self._enviar_directo(
                "Toma nota: (dime aqui lo que quieres apuntar)")),
            ("Mis rutinas", lambda: self._enviar_directo(
                "Que rutinas tengo activas?")),
            ("Nuevo chat", self._nuevo_chat),
        ]:
            tk.Button(left, text=txt, command=cmd,
                bg=C["bg3"], fg=C["fg"], relief="flat",
                font=("Segoe UI", 9), cursor="hand2",
                pady=4).pack(fill="x", padx=20, pady=2)

        # --- Panel central: CHAT ---
        center = tk.Frame(body, bg=C["bg"])
        center.pack(side="left", fill="both", expand=True)

        # Chat
        self.chat = scrolledtext.ScrolledText(center,
            bg=C["bg2"], fg=C["fg"], insertbackground=C["fg"],
            font=("Segoe UI", 10), wrap="word", relief="flat",
            padx=14, pady=14)
        self.chat.pack(fill="both", expand=True, padx=12, pady=12)
        self.chat.tag_config("user", foreground=C["user"],
            font=("Segoe UI", 10, "bold"))
        self.chat.tag_config("jarvis", foreground=C["jarvis"],
            font=("Segoe UI", 10))
        self.chat.tag_config("sys", foreground=C["sys"],
            font=("Segoe UI", 9, "italic"))
        self.chat.tag_config("tool", foreground=C["tool"],
            font=("Consolas", 9))
        self.chat.tag_config("err", foreground=C["err"])
        self.chat.tag_config("warn", foreground=C["warn"])

        # Input
        input_frame = tk.Frame(center, bg=C["bg"])
        input_frame.pack(fill="x", padx=12, pady=(0, 12))

        # Boton adjuntar imagen (para analisis con vision)
        self.imagen_adjunta = None
        self.btn_adjuntar = tk.Button(input_frame, text="[+img]",
            command=self._adjuntar_imagen,
            bg=C["bg3"], fg=C["fg"], relief="flat",
            font=("Segoe UI", 10), cursor="hand2",
            padx=10, pady=8)
        self.btn_adjuntar.pack(side="left", padx=(0, 4), ipady=2)

        # Boton captura de pantalla (analiza directamente)
        tk.Button(input_frame, text="[camara]",
            command=self._adjuntar_pantalla,
            bg=C["bg3"], fg=C["fg"], relief="flat",
            font=("Segoe UI", 10), cursor="hand2",
            padx=10, pady=8).pack(side="left", padx=(0, 4), ipady=2)

        self.entry = tk.Entry(input_frame,
            bg=C["bg2"], fg=C["fg"], insertbackground=C["fg"],
            relief="flat", font=("Segoe UI", 11))
        self.entry.pack(side="left", fill="x", expand=True, ipady=10)
        self.entry.bind("<Return>", lambda e: self._enviar())

        tk.Button(input_frame, text="ENVIAR", command=self._enviar,
            bg=C["jarvis"], fg="white", relief="flat",
            font=("Segoe UI", 10, "bold"), cursor="hand2",
            padx=20).pack(side="left", padx=6, ipady=8)

        # --- Panel derecho: VISOR MULTIMEDIA (colapsable) ---
        self.visor_frame = tk.Frame(body, bg=C["bg2"], width=380)
        self.visor_frame.pack(side="right", fill="y")
        self.visor_frame.pack_propagate(False)
        self._construir_visor()

    # ============================================================
    # VISOR MULTIMEDIA (imagenes, videos-mini, docs generados)
    # ============================================================
    def _construir_visor(self):
        """Construye el panel visor multimedia a la derecha."""
        # Header del visor
        vh = tk.Frame(self.visor_frame, bg=C["bg3"], height=44)
        vh.pack(fill="x")
        vh.pack_propagate(False)

        tk.Label(vh, text="VISOR MULTIMEDIA",
            bg=C["bg3"], fg=C["jarvis"],
            font=("Segoe UI", 10, "bold")).pack(side="left", padx=12, pady=10)

        tk.Button(vh, text="Limpiar",
            command=self._limpiar_visor,
            bg=C["bg2"], fg=C["fg2"], relief="flat",
            font=("Segoe UI", 8), cursor="hand2",
            padx=8, pady=2).pack(side="right", padx=8, pady=10)

        # Scrollable content area
        cont = tk.Frame(self.visor_frame, bg=C["bg2"])
        cont.pack(fill="both", expand=True)

        self.visor_canvas = tk.Canvas(cont, bg=C["bg2"],
                                        highlightthickness=0)
        sb = tk.Scrollbar(cont, orient="vertical",
                           command=self.visor_canvas.yview,
                           bg=C["bg3"])
        self.visor_content = tk.Frame(self.visor_canvas, bg=C["bg2"])

        self.visor_content.bind("<Configure>",
            lambda e: self.visor_canvas.configure(
                scrollregion=self.visor_canvas.bbox("all")))
        self.visor_canvas.create_window((0, 0),
            window=self.visor_content, anchor="nw", width=360)
        self.visor_canvas.configure(yscrollcommand=sb.set)
        self.visor_canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        # Scroll con rueda
        def _mw(e):
            self.visor_canvas.yview_scroll(-int(e.delta / 60), "units")
        self.visor_canvas.bind("<Enter>",
            lambda e: self.visor_canvas.bind_all("<MouseWheel>", _mw))
        self.visor_canvas.bind("<Leave>",
            lambda e: self.visor_canvas.unbind_all("<MouseWheel>"))

        # Mensaje inicial
        self._mensaje_vacio_visor()

        # Guardar referencias a imagenes para evitar GC
        self._visor_imgs = []

    def _mensaje_vacio_visor(self):
        for w in self.visor_content.winfo_children():
            w.destroy()
        tk.Label(self.visor_content,
            text="El visor esta vacio.\n\n"
                 "Pideme cosas como:\n"
                 "  * genera una imagen de...\n"
                 "  * crea un video de 3 escenas sobre...\n"
                 "  * crea un documento Word sobre...\n"
                 "  * anade el habito de correr\n\n"
                 "Los resultados apareceran aqui.",
            bg=C["bg2"], fg=C["fg2"],
            font=("Segoe UI", 9),
            justify="left").pack(padx=16, pady=30, anchor="w")

    def _limpiar_visor(self):
        self._visor_imgs = []
        self._mensaje_vacio_visor()

    def _visor_add_titulo(self, texto, color=None):
        """Agrega un titulo/separador al visor."""
        # Si es el primer item, borrar mensaje vacio
        if len(self.visor_content.winfo_children()) == 1:
            first = self.visor_content.winfo_children()[0]
            if isinstance(first, tk.Label) and "vacio" in first.cget("text"):
                first.destroy()
        header = tk.Frame(self.visor_content, bg=C["bg2"])
        header.pack(fill="x", padx=8, pady=(12, 4))
        tk.Label(header, text=texto,
            bg=C["bg2"], fg=color or C["jarvis"],
            font=("Segoe UI", 10, "bold"),
            anchor="w").pack(fill="x")
        # linea separadora
        tk.Frame(header, bg=color or C["jarvis"], height=2).pack(fill="x", pady=(2, 0))

    def _visor_add_imagen(self, ruta_img, subtitulo=""):
        """Agrega una imagen al visor."""
        try:
            from PIL import Image, ImageTk
            img = Image.open(ruta_img)
            # Redimensionar a max 340px de ancho
            w, h = img.size
            if w > 340:
                ratio = 340 / w
                img = img.resize((340, int(h * ratio)))
            photo = ImageTk.PhotoImage(img)
            self._visor_imgs.append(photo)  # evitar GC

            card = tk.Frame(self.visor_content, bg=C["bg3"])
            card.pack(fill="x", padx=8, pady=4)

            lbl = tk.Label(card, image=photo, bg=C["bg3"], cursor="hand2")
            lbl.pack(padx=6, pady=6)
            # Click abre imagen completa
            lbl.bind("<Button-1>",
                lambda e, r=ruta_img: self._abrir_visor_completo(r))

            if subtitulo:
                tk.Label(card, text=subtitulo,
                    bg=C["bg3"], fg=C["fg"],
                    font=("Segoe UI", 9),
                    wraplength=340, justify="left").pack(padx=6, pady=(0, 6))
            return True
        except Exception as e:
            tk.Label(self.visor_content,
                text=f"[imagen: {os.path.basename(ruta_img)}] (error: {e})",
                bg=C["bg2"], fg=C["warn"],
                font=("Segoe UI", 8)).pack(padx=8, pady=2, anchor="w")
            return False

    def _visor_add_texto(self, texto, color=None, mono=False):
        """Agrega texto/parrafo al visor."""
        tk.Label(self.visor_content, text=texto,
            bg=C["bg2"], fg=color or C["fg"],
            font=("Consolas" if mono else "Segoe UI",
                   9, "italic" if not mono else "normal"),
            wraplength=340, justify="left", anchor="w"
            ).pack(fill="x", padx=12, pady=2)

    def _visor_add_boton(self, texto, comando, color=None):
        """Agrega un boton (ej: 'Abrir en Studio')."""
        tk.Button(self.visor_content, text=texto, command=comando,
            bg=color or C["jarvis"], fg="white", relief="flat",
            font=("Segoe UI", 9, "bold"), cursor="hand2",
            padx=12, pady=4).pack(padx=8, pady=(4, 10))

    def _abrir_visor_completo(self, ruta_img):
        """Abre una imagen a tamano grande en una ventana emergente."""
        try:
            from PIL import Image, ImageTk
            win = tk.Toplevel(self.root)
            win.title(os.path.basename(ruta_img))
            win.configure(bg=C["bg"])
            img = Image.open(ruta_img)
            w, h = img.size
            max_w, max_h = 1000, 800
            if w > max_w or h > max_h:
                ratio = min(max_w / w, max_h / h)
                img = img.resize((int(w * ratio), int(h * ratio)))
            photo = ImageTk.PhotoImage(img)
            lbl = tk.Label(win, image=photo, bg=C["bg"])
            lbl.image = photo  # evitar GC
            lbl.pack(padx=10, pady=10)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir: {e}")

    def _loop_animacion(self):
        self.esfera.animar()
        self.root.after(50, self._loop_animacion)

    def _set_estado(self, estado, texto):
        self.esfera.set_estado(estado)
        self.lbl_estado.config(text=texto)

    def _mensaje_bienvenida(self):
        nombre = ""
        for h in self.memoria.get("hechos", []):
            if "me llamo" in h["texto"].lower() or "mi nombre" in h["texto"].lower():
                m = re.search(r"(?:me llamo|mi nombre es)\s+(\w+)",
                               h["texto"], re.IGNORECASE)
                if m:
                    nombre = m.group(1)
                    break
        saludo = f"Hola{' ' + nombre if nombre else ''}, soy Jarvis."
        info = (f"Recuerdo {len(self.memoria.get('hechos', []))} hechos "
                 f"sobre ti y {len(self.memoria.get('conversaciones', []))} conversaciones.")
        tools_info = ("Puedo ejecutar herramientas: abrir webs, buscar en Google, "
                       "generar imagenes, decirte la hora y el clima, hacer capturas, "
                       "abrir apps NOVA, calcular, crear recordatorios y mucho mas. "
                       "Solo pidemelo con lenguaje natural.")
        self._chat_msg("Jarvis", f"{saludo}\n{info}\n\n{tools_info}", "jarvis")

    def _chat_msg(self, autor, texto, tag):
        prefijo = f"{autor}: " if autor else ""
        self.chat.insert("end", f"\n{prefijo}", tag)
        self.chat.insert("end", f"{texto}\n", tag)
        self.chat.see("end")

    def _enviar(self):
        texto = self.entry.get().strip()
        if not texto: return
        self.entry.delete(0, "end")
        self._enviar_directo(texto)

    def _enviar_directo(self, texto):
        if self.procesando:
            self._chat_msg("Sistema", "Espera, estoy procesando la anterior...", "sys")
            return
        # Comandos con /
        if texto.startswith("/"):
            self._comando(texto)
            return
        self._chat_msg("Tu", texto, "user")
        self._procesar(texto)

    def _comando(self, cmd):
        c = cmd.lower().strip()
        if c in ("/nuevo", "/reset", "/new"):
            self._nuevo_chat()
        elif c in ("/memoria", "/mem"):
            self._ver_memoria()
        elif c in ("/tools", "/herramientas"):
            self._ver_herramientas()
        elif c in ("/ayuda", "/help", "/?"):
            self._ver_ayuda()
        elif c.startswith("/recordar "):
            hecho = cmd[10:]
            r, ok = self.herramientas.recordar_hecho(hecho)
            self._chat_msg("Jarvis", r, "jarvis" if ok else "err")
        elif c == "/limpiar":
            self.chat.delete("1.0", "end")
            self._mensaje_bienvenida()
        else:
            self._chat_msg("Sistema",
                f"Comando desconocido: {cmd}\n"
                "Prueba: /nuevo /memoria /tools /ayuda /recordar <hecho> /limpiar",
                "sys")

    def _nuevo_chat(self):
        self.historial_actual = []
        self.chat.delete("1.0", "end")
        self._mensaje_bienvenida()
        self._set_estado("idle", "Listo")
        self.imagen_adjunta = None
        self.btn_adjuntar.config(text="[+img]", bg=C["bg3"])

    def _adjuntar_imagen(self):
        """Deja que el usuario elija una imagen para analizar con vision."""
        ruta = filedialog.askopenfilename(
            title="Elige una imagen para adjuntar",
            filetypes=[("Imagenes", "*.png *.jpg *.jpeg *.gif *.webp *.bmp"),
                        ("Todo", "*.*")])
        if not ruta: return
        self.imagen_adjunta = ruta
        self.btn_adjuntar.config(text="[img OK]", bg=C["ok"])
        # Mostrar preview en el visor
        self._visor_add_titulo("IMAGEN ADJUNTA", C["ok"])
        self._visor_add_imagen(ruta,
            "Escribe algo en el chat para preguntar sobre esta imagen")
        self._chat_msg("Sistema",
            f"Imagen adjuntada: {os.path.basename(ruta)}. Escribe tu pregunta.",
            "sys")

    def _adjuntar_pantalla(self):
        """Captura la pantalla y la adjunta al chat."""
        if not PIL_OK:
            messagebox.showwarning("Sin Pillow", "Instala Pillow: pip install Pillow")
            return
        try:
            # Ocultar la ventana brevemente para no capturarla
            self.root.iconify()
            self.root.after(500, lambda: self._hacer_captura_adjunta())
        except Exception as e:
            self._chat_msg("Sistema", f"Error capturando: {e}", "err")

    def _hacer_captura_adjunta(self):
        try:
            img = ImageGrab.grab()
            ruta = os.path.join(CAPTURAS,
                f"screen_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
            img.save(ruta)
            self.root.deiconify()
            self.imagen_adjunta = ruta
            self.btn_adjuntar.config(text="[img OK]", bg=C["ok"])
            self._visor_add_titulo("PANTALLA CAPTURADA", C["ok"])
            self._visor_add_imagen(ruta, "Escribe tu pregunta en el chat")
            self._chat_msg("Sistema",
                "Captura adjuntada. Pregunta lo que quieras sobre lo que ves en pantalla.",
                "sys")
        except Exception as e:
            self.root.deiconify()
            self._chat_msg("Sistema", f"Error: {e}", "err")

    def _procesar(self, texto):
        if not self.api_key and self.proveedor != "Ollama":
            self._chat_msg("Sistema", "Configura tu API Key en Config IA.", "err")
            return

        self.procesando = True
        self._set_estado("think", "Pensando...")

        # Autodeteccion de hechos del usuario (nombre, edad, ciudad, etc.)
        hechos_detectados = extraer_hechos(texto)
        for h in hechos_detectados:
            # Guardar en perfil del brain compartido
            self.brain.perfil(h["campo"], h["valor"])

        # Capturar imagen adjunta antes de la tarea
        img_para_enviar = self.imagen_adjunta

        def tarea():
            # Construir prompt de sistema con memoria y tools
            system = self._construir_system_prompt()
            # Añadir mensaje al historial
            self.historial_actual.append({"role": "user", "content": texto})
            # Llamar IA (con imagen si esta adjunta)
            resp, err = llamar_llm(self.proveedor, self.modelo, self.api_key,
                                    system, texto,
                                    historial=self.historial_actual[:-1],
                                    imagen=img_para_enviar)
            if err:
                self.root.after(0, lambda: self._chat_msg(
                    "Jarvis", f"Error: {err}", "err"))
                self.root.after(0, lambda: self._set_estado("idle", "Listo"))
                self.procesando = False
                return

            # Procesar la respuesta (puede incluir tools)
            self.root.after(0, self._responder, resp, texto)
            # Limpiar imagen adjunta despues de usarla
            if img_para_enviar:
                self.root.after(0, self._limpiar_adjunto)

        threading.Thread(target=tarea, daemon=True).start()

    def _limpiar_adjunto(self):
        """Limpia la imagen adjunta despues de enviarla."""
        self.imagen_adjunta = None
        self.btn_adjuntar.config(text="[+img]", bg=C["bg3"])

    def _construir_system_prompt(self):
        """Construye el prompt de sistema con memoria y descripcion de tools."""
        # Combina hechos locales + hechos del cerebro compartido
        hechos_str = ""
        hechos_locales = self.memoria.get("hechos", [])
        hechos_brain = self.brain.hechos_recientes(30)
        # Unir sin duplicar
        vistos = set()
        todos = []
        for h in hechos_locales[-20:] + hechos_brain:
            if h["texto"] not in vistos:
                vistos.add(h["texto"])
                todos.append(h)
        if todos:
            hechos_str = "HECHOS QUE RECUERDO SOBRE EL USUARIO:\n"
            for h in todos[-30:]:
                mod = h.get("modulo", "")
                marca = f" [via {mod}]" if mod and mod != "jarvis" else ""
                hechos_str += f"  - {h['texto']}{marca}\n"

        # Perfil estructurado del brain
        perfil = self.brain.perfil()
        perfil_str = ""
        if perfil:
            perfil_str = "\nPERFIL DEL USUARIO:\n"
            for k, v in perfil.items():
                if v:
                    perfil_str += f"  {k}: {v}\n"

        # Ultimas conversaciones para contexto
        convs_str = ""
        convs = self.memoria.get("conversaciones", [])
        if convs:
            convs_str = "\nCONVERSACIONES ANTERIORES (contexto):\n"
            for c in convs[-3:]:
                convs_str += f"  Usuario: {c.get('user','')[:100]}\n"
                convs_str += f"  Tu: {c.get('nova','')[:100]}\n"

        # Actividad reciente del ecosistema NOVA
        eventos = self.brain.eventos_recientes(5)
        eventos_str = ""
        if eventos:
            eventos_str = "\nACTIVIDAD RECIENTE EN OTROS MODULOS NOVA:\n"
            for e in eventos:
                eventos_str += f"  - {e['modulo']}: {e['tipo']}\n"

        tools_str = ""
        if self.tools_habilitadas:
            tools_str = """

HERRAMIENTAS QUE PUEDES USAR:
Para usar una herramienta, incluye en tu respuesta la etiqueta:
  [[TOOL:nombre|arg1|arg2]]

Herramientas disponibles:
  [[TOOL:abrir_web|url]]              - abre una URL en el navegador
  [[TOOL:buscar_web|query]]           - busca en Google
  [[TOOL:buscar_youtube|query]]       - busca en YouTube
  [[TOOL:buscar_wikipedia|termino]]   - resume de Wikipedia
  [[TOOL:abrir_app|nombre]]           - abre una app (calc, notepad, chrome...)
  [[TOOL:listar_carpeta|ruta]]        - lista archivos de una carpeta
  [[TOOL:leer_archivo|ruta]]          - lee el contenido de un archivo
  [[TOOL:escribir_archivo|ruta|contenido]] - crea/sobreescribe un archivo
  [[TOOL:ejecutar_comando|cmd]]       - ejecuta comando shell (pide confirmacion)
  [[TOOL:captura_pantalla]]           - captura la pantalla
  [[TOOL:generar_imagen|prompt_en_ingles]] - genera imagen (aparece EN el visor)
  [[TOOL:crear_video_mini|tema|n_escenas]] - mini-video de N imagenes + narracion EN el visor
  [[TOOL:decir_hora]]                 - dice la hora actual
  [[TOOL:decir_clima|ciudad]]         - dice el clima
  [[TOOL:calcular_expresion|expr]]    - calcula matematicas (ej: 2+2*5)
  [[TOOL:recordar_hecho|texto]]       - guarda un hecho en memoria
  [[TOOL:olvidar_hecho|indice]]       - olvida un hecho por indice
  [[TOOL:listar_hechos]]              - lista lo que recuerdas
  [[TOOL:crear_recordatorio|texto|minutos]] - crea recordatorio
  [[TOOL:listar_recordatorios]]       - lista recordatorios activos
  [[TOOL:abrir_modulo_nova|nombre]]   - abre otro modulo NOVA (reader, studio, coach, office...)

CREACION DE DOCUMENTOS (via NOVA Office):
  [[TOOL:crear_word|tema|tipo]]       - crea documento Word (tipo: informe/ensayo/carta/cuento...)
  [[TOOL:crear_excel|tema]]           - crea hoja Excel con datos sobre el tema

CONTROL DE NOVA COACH:
  [[TOOL:add_habito_coach|nombre|emoji]] - agrega un habito nuevo
  [[TOOL:crear_meta_coach|texto]]     - agrega una meta

LECTURA DE PDFs:
  [[TOOL:resumir_pdf|ruta]]           - lee un PDF y resume en el visor

VISION (analisis de imagenes con IA multimodal):
  [[TOOL:analizar_pantalla|pregunta]]    - captura pantalla y la analiza
  [[TOOL:analizar_imagen|ruta|pregunta]] - analiza imagen local
  [[TOOL:extraer_texto_imagen|ruta]]     - extrae texto de una imagen (OCR)

MODO AGENTE (planificacion multi-paso):
  [[TOOL:modo_agente|objetivo]]       - la IA planifica y ejecuta varios pasos
      Ejemplo: [[TOOL:modo_agente|investiga la historia de Python, crea un word con el resumen y una imagen ilustrativa]]

AUTOMATIZACIONES / RUTINAS PROGRAMADAS:
  [[TOOL:crear_rutina|nombre|HH:MM|acciones]] - crea rutina diaria a hora fija
  [[TOOL:listar_rutinas]]             - lista rutinas activas
  [[TOOL:borrar_rutina|nombre]]       - elimina una rutina

RAG (chat con TUS documentos):
  [[TOOL:rag_indexar|ruta_carpeta]]   - indexa PDFs/txt/docx de una carpeta
  [[TOOL:rag_preguntar|consulta]]     - pregunta usando tus documentos indexados
  [[TOOL:rag_resumen]]                - ver que hay indexado
  [[TOOL:rag_limpiar]]                - borra el indice

CONTROL DE MUSICA:
  [[TOOL:musica_play_pause]]          - pausa/reanuda musica en cualquier reproductor
  [[TOOL:musica_siguiente]]           - siguiente cancion
  [[TOOL:musica_anterior]]            - cancion anterior
  [[TOOL:musica_volumen|subir|3]]     - sube/baja/mute volumen
  [[TOOL:musica_reproducir|query]]    - busca cancion en YouTube Music
  [[TOOL:musica_mood|mood]]           - reproduce por mood: relax, concentracion, energia, fiesta, chill, etc.
  [[TOOL:musica_spotify|query]]       - busca en Spotify

COMUNICACION:
  [[TOOL:abrir_whatsapp_web|numero+mensaje]] - abre chat WhatsApp (numero+mensaje opcional)
  [[TOOL:abrir_gmail_web|destino|asunto|cuerpo]] - compone email en Gmail
  [[TOOL:abrir_correo|destino|asunto|cuerpo]] - compone en app de correo local

NOTAS RAPIDAS:
  [[TOOL:nota_rapida|texto]]          - guarda una nota rapida
  [[TOOL:leer_notas|10]]              - lee las ultimas N notas

CALENDARIO/MAPS:
  [[TOOL:abrir_calendar|evento]]      - abre Google Calendar (con evento opcional)
  [[TOOL:buscar_maps|lugar]]          - busca en Google Maps

DEV:
  [[TOOL:buscar_github|query]]        - busca repos GitHub
  [[TOOL:abrir_stackoverflow|pregunta]] - busca en Stack Overflow

VISION CONTINUA:
  [[TOOL:vision_continua_iniciar|10]] - activa modo vision continua (cada N seg)
  [[TOOL:vision_continua_parar]]      - desactiva vision continua
  [[TOOL:que_veo_ahora|pregunta]]     - captura pantalla YA y responde a la pregunta

REGLAS:
1. Si el usuario pide algo que una herramienta puede hacer, USA la herramienta.
2. Puedes usar VARIAS herramientas en una respuesta.
3. Cuando uses herramientas, EXPLICA brevemente que estas haciendo.
4. NO uses etiquetas TOOL para cosas que no lo requieren (charlas casuales).
5. Cuando el usuario te diga algo personal importante (nombre, edad, gustos, cumple),
   USA [[TOOL:recordar_hecho|...]] automaticamente.
6. IMPORTANTE: las imagenes y videos SIEMPRE se muestran en el visor derecho,
   NO se abren archivos externos.
7. Si el usuario pide "crea un video de X" usa crear_video_mini.
8. Si pide "hazme un word/excel de X" usa crear_word/crear_excel.
9. Si te adjuntan una imagen (aparece imagen en el mensaje), respondes describiendo
   lo que ves con detalle. No necesitas herramientas para eso.
10. Si el objetivo requiere VARIAS acciones en cadena (investigar Y crear Y guardar),
    usa [[TOOL:modo_agente|objetivo]] para que se planifique automaticamente."""

        # === PROMPT SEGUN LA PERSONALIDAD ACTIVA ===
        if PERS_OK:
            prompt_base = obtener_prompt(self.personalidad)
        else:
            prompt_base = ("Eres Jarvis, un asistente personal poderoso estilo Iron Man. "
                            "Hablas en espanol, eres eficiente, respetuoso y con toque elegante. "
                            "Respondes de forma concisa pero completa.")

        # === CONTEXTO DE VISION CONTINUA (si esta activa) ===
        vision_str = ""
        if self.vision_continua_activa and self.vision_continua_ultimo_analisis:
            vision_str = (f"\nCONTEXTO ACTUAL DE PANTALLA (vision continua activa):\n"
                           f"  {self.vision_continua_ultimo_analisis}\n"
                           f"  (Sabes que esta viendo el usuario en su pantalla)\n")

        base = (f"{prompt_base}\n\n"
                 f"{perfil_str}{hechos_str}{convs_str}{eventos_str}{vision_str}{tools_str}")
        return base

    def _responder(self, resp, pregunta_original):
        # Extraer tools de la respuesta
        tools = parsear_tools(resp) if self.tools_habilitadas else []

        # Texto sin las etiquetas TOOL
        texto_limpio = TOOL_RE.sub("", resp).strip()

        # Mostrar la respuesta limpia
        if texto_limpio:
            self._chat_msg("Jarvis", texto_limpio, "jarvis")

        # Ejecutar cada herramienta
        for nombre, args, _ in tools:
            self._set_estado("tool", f"Usando {nombre}...")
            self._chat_msg("", f"  >> Ejecutando {nombre}({', '.join(args)[:80]})", "tool")
            self._ejecutar_tool(nombre, args)

        # Guardar en historial
        self.historial_actual.append({"role": "assistant", "content": resp})
        self.memoria.setdefault("conversaciones", []).append({
            "fecha": datetime.now().isoformat(),
            "user": pregunta_original,
            "nova": texto_limpio,
        })
        # Limitar a 200 conversaciones
        if len(self.memoria["conversaciones"]) > 200:
            self.memoria["conversaciones"] = self.memoria["conversaciones"][-200:]
        guardar_json(MEMORIA, self.memoria)

        # Hablar si activado
        if self.voz_activada and texto_limpio:
            self._set_estado("speak", "Hablando...")
            def hablar():
                hablar_voz(texto_limpio[:500], self.voz)
                self.root.after(0, lambda: self._set_estado("idle", "Listo"))
                # Si estamos en modo continuo, seguir escuchando
                if self.modo_continuo:
                    self.root.after(500, self._escuchar_una)
            threading.Thread(target=hablar, daemon=True).start()
        else:
            self._set_estado("idle", "Listo")
            if self.modo_continuo:
                self.root.after(500, self._escuchar_una)

        self.procesando = False

    def _ejecutar_tool(self, nombre, args):
        # Buscar el metodo en Herramientas
        metodo = getattr(self.herramientas, nombre, None)
        if not metodo:
            self._chat_msg("", f"  << Herramienta desconocida: {nombre}", "err")
            return
        try:
            resultado, ok = metodo(*args) if args else metodo()
            resultado_corto = resultado[:400] + "..." if len(resultado) > 400 else resultado
            tag = "tool" if ok else "warn"
            self._chat_msg("", f"  << {resultado_corto}", tag)
        except TypeError as e:
            self._chat_msg("", f"  << Error argumentos {nombre}: {e}", "err")
        except Exception as e:
            self._chat_msg("", f"  << Error ejecutando {nombre}: {e}", "err")

    # ============================================================
    # VOZ
    # ============================================================
    def _toggle_voz(self):
        self.voz_activada = not self.voz_activada
        self.cfg["jarvis_voz_activada"] = self.voz_activada
        guardar_json(CONFIG, self.cfg)
        self.btn_voz.config(text="Voz: ON" if self.voz_activada else "Voz: OFF",
            bg=C["ok"] if self.voz_activada else C["bg3"])

    def _toggle_tools(self):
        self.tools_habilitadas = not self.tools_habilitadas
        self.cfg["jarvis_tools"] = self.tools_habilitadas
        guardar_json(CONFIG, self.cfg)
        self.btn_tools.config(
            text="Herramientas: " + ("ON" if self.tools_habilitadas else "OFF"),
            bg=C["tool"] if self.tools_habilitadas else C["bg3"])

    def _toggle_wake_word(self):
        """Activa/desactiva el sistema de wake word 'Hey Jarvis'."""
        if not VOZ_PRO_OK:
            messagebox.showwarning("Voz profesional",
                "Instala:\n  pip install SpeechRecognition PyAudio\n\n"
                "Whisper (opcional, mas preciso):\n  pip install openai-whisper")
            return

        if self.voz_pro_activa:
            # Desactivar
            if self.voz_pro:
                self.voz_pro.parar()
            self.voz_pro_activa = False
            self.btn_wake.config(text="Wake 'Hey Jarvis': OFF",
                                   bg=C["bg3"])
            self._set_estado("idle", "Listo")
            return

        # Activar
        try:
            self.voz_pro = VozProfesional(
                wake_words=["hey jarvis", "jarvis", "oye jarvis"],
                on_wake=self._wake_activado,
                on_transcripcion=self._wake_transcripcion,
                on_estado=self._wake_estado,
            )
            if self.voz_pro.iniciar():
                self.voz_pro_activa = True
                self.btn_wake.config(text="Wake 'Hey Jarvis': ON",
                                       bg=C["ok"])
                self._chat_msg("Sistema",
                    "Wake word activada. Di 'Hey Jarvis' para hablar.",
                    "sys")
            else:
                messagebox.showerror("Error", "No se pudo iniciar la voz profesional")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _wake_activado(self):
        """Callback cuando se detecta 'Hey Jarvis' sin comando adicional."""
        self.root.after(0, lambda: self._set_estado("listen", "Te escucho..."))

    def _wake_transcripcion(self, texto):
        """Callback cuando llega un comando completo por voz."""
        self.root.after(0, lambda: self._enviar_directo(texto))

    def _wake_estado(self, texto):
        """Callback de cambios de estado de la voz."""
        # Solo actualizar si no estamos procesando
        if not self.procesando:
            self.root.after(0, lambda t=texto: self._set_estado("idle", t))

    def _toggle_continuo(self):
        if not SR_OK:
            messagebox.showwarning("Sin reconocimiento de voz",
                "Instala:\n  pip install SpeechRecognition PyAudio")
            return
        self.modo_continuo = not self.modo_continuo
        if self.modo_continuo:
            self.btn_continuo.config(text="Parar escucha continua",
                bg=C["err"])
            self._escuchar_una()
        else:
            self.btn_continuo.config(text="Modo escucha continua",
                bg=C["bg3"])

    def _escuchar_una(self):
        if not SR_OK:
            messagebox.showwarning("Sin reconocimiento de voz",
                "Instala:\n  pip install SpeechRecognition PyAudio")
            return
        if self.procesando:
            return
        self._set_estado("listen", "Escuchando...")

        def tarea():
            try:
                r = sr.Recognizer()
                with sr.Microphone() as source:
                    r.adjust_for_ambient_noise(source, duration=0.5)
                    audio = r.listen(source, timeout=6, phrase_time_limit=15)
                self.root.after(0, lambda: self._set_estado("think", "Reconociendo..."))
                texto = r.recognize_google(audio, language="es-ES")
                self.root.after(0, lambda: self._enviar_directo(texto))
            except sr.WaitTimeoutError:
                self.root.after(0, lambda: self._set_estado("idle", "Silencio - Listo"))
                if self.modo_continuo:
                    self.root.after(500, self._escuchar_una)
            except Exception as e:
                self.root.after(0, lambda: self._chat_msg(
                    "Sistema", f"Error voz: {e}", "err"))
                self.root.after(0, lambda: self._set_estado("idle", "Listo"))

        threading.Thread(target=tarea, daemon=True).start()

    # ============================================================
    # RECORDATORIOS
    # ============================================================
    def _loop_vision_continua(self):
        """Loop de vision continua: cada X seg captura y analiza pantalla."""
        if not self.vision_continua_activa:
            return
        if not PIL_OK or not self.api_key:
            self.vision_continua_activa = False
            return
        # Capturar y analizar en background
        def tarea():
            try:
                # Ocultar Jarvis brevemente
                try: self.root.attributes("-alpha", 0.0)
                except: pass
                time.sleep(0.2)
                img = ImageGrab.grab()
                try: self.root.attributes("-alpha", 1.0)
                except: pass
                # Guardar la ultima captura
                self.vision_continua_ultima_captura = img
                # Preguntar a la IA de forma muy breve
                system = ("Describe MUY brevemente (max 30 palabras) que aplicacion "
                           "o contenido se ve en la pantalla. Solo el hecho, sin "
                           "opinion.")
                resp, err = llamar_llm(self.proveedor, self.modelo, self.api_key,
                    system, "Que ves en pantalla?", timeout=45, imagen=img)
                if not err and resp:
                    self.vision_continua_ultimo_analisis = resp.strip()
                    # Guardar en Brain como hecho reciente
                    if hasattr(self, "brain"):
                        self.brain.stat("jarvis", "vision_capturas", incrementar=1)
            except Exception as e:
                print(f"[VisionContinua] Error: {e}")
        threading.Thread(target=tarea, daemon=True).start()
        # Programar siguiente iteracion
        self.root.after(self.vision_continua_intervalo * 1000,
            self._loop_vision_continua)

    def _check_recordatorios(self):
        ahora = datetime.now()
        for r in self.recordatorios:
            if r.get("notificado"):
                continue
            try:
                cuando = datetime.fromisoformat(r["hora"])
                if cuando <= ahora:
                    r["notificado"] = True
                    self._chat_msg("JARVIS!!",
                        f"RECORDATORIO: {r['texto']}", "warn")
                    if self.voz_activada:
                        threading.Thread(
                            target=lambda: hablar_voz(
                                f"Recordatorio: {r['texto']}", self.voz),
                            daemon=True).start()
                    guardar_json(RECORDATORIOS, self.recordatorios)
            except Exception:
                pass
        self.root.after(30000, self._check_recordatorios)

    def _check_rutinas(self):
        """Comprueba cada minuto si toca ejecutar alguna rutina."""
        ahora = datetime.now()
        hora_actual = ahora.strftime("%H:%M")
        clave_hoy = ahora.strftime("%Y-%m-%d")
        for r in self.rutinas:
            if not r.get("activa", True):
                continue
            hora_r = r.get("hora", "").strip()
            # Formato "HH:MM" diario
            if re.match(r"^\d{1,2}:\d{2}$", hora_r):
                # Normalizar formato HH:MM
                partes = hora_r.split(":")
                hora_r_norm = f"{int(partes[0]):02d}:{partes[1]}"
                if hora_r_norm == hora_actual:
                    # Comprobar que no la hemos ejecutado hoy ya
                    if r.get("ultima_ejecucion", "").startswith(clave_hoy):
                        continue
                    self._ejecutar_rutina(r)
        self.root.after(60000, self._check_rutinas)  # cada minuto

    def _ejecutar_rutina(self, rutina):
        """Ejecuta las acciones de una rutina como si el usuario las pidiera."""
        rutina["ultima_ejecucion"] = datetime.now().isoformat()
        guardar_json(RUTINAS, self.rutinas)
        self._chat_msg("RUTINA",
            f"Ejecutando rutina '{rutina['nombre']}': {rutina['acciones']}",
            "warn")
        # Enviar las acciones como si fueran pregunta del usuario
        self.root.after(500, lambda: self._procesar(rutina["acciones"]))

    # ============================================================
    # DIALOGOS
    # ============================================================
    def _ver_memoria(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Memoria de Jarvis")
        dlg.geometry("640x520")
        dlg.configure(bg=C["bg"])

        tk.Label(dlg, text="MEMORIA DE JARVIS",
            bg=C["bg"], fg=C["jarvis"],
            font=("Segoe UI", 14, "bold")).pack(pady=12)

        hechos = self.memoria.get("hechos", [])
        tk.Label(dlg, text=f"Hechos guardados: {len(hechos)}",
            bg=C["bg"], fg=C["fg"],
            font=("Segoe UI", 10)).pack()

        txt = scrolledtext.ScrolledText(dlg, bg=C["bg2"], fg=C["fg"],
            font=("Consolas", 10), wrap="word", relief="flat",
            padx=10, pady=10)
        txt.pack(fill="both", expand=True, padx=12, pady=8)

        if hechos:
            for i, h in enumerate(hechos):
                txt.insert("end", f"[{i}] {h['texto']}\n     ({h.get('fecha','')[:16]})\n\n")
        else:
            txt.insert("end", "Aun no recuerdo nada sobre ti.\n"
                              "Dime cosas como 'me llamo X', 'me gusta Y', 'mi cumple es Z'.")

        bf = tk.Frame(dlg, bg=C["bg"])
        bf.pack(pady=10)
        def borrar_todo():
            if messagebox.askyesno("Borrar todo?",
                "Esto borra TODOS los hechos y conversaciones. Seguro?"):
                self.memoria = {"hechos": [], "conversaciones": []}
                guardar_json(MEMORIA, self.memoria)
                dlg.destroy()
        tk.Button(bf, text="Borrar TODA la memoria", command=borrar_todo,
            bg=C["err"], fg="white", relief="flat",
            font=("Segoe UI", 10, "bold"), padx=16, pady=6,
            cursor="hand2").pack(side="left", padx=6)
        tk.Button(bf, text="Cerrar", command=dlg.destroy,
            bg=C["bg3"], fg=C["fg"], relief="flat",
            font=("Segoe UI", 10), padx=16, pady=6).pack(side="left", padx=6)

    def _ver_herramientas(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Herramientas disponibles")
        dlg.geometry("700x600")
        dlg.configure(bg=C["bg"])

        tk.Label(dlg, text="HERRAMIENTAS DE JARVIS",
            bg=C["bg"], fg=C["tool"],
            font=("Segoe UI", 14, "bold")).pack(pady=12)

        info = """Jarvis puede ejecutar estas herramientas cuando se lo pidas
en lenguaje natural. NO necesitas usar la sintaxis tecnica, solo pide
lo que quieras y Jarvis decidira que herramienta usar.

WEB
  - Abrir URL         "abre google.com"
  - Buscar en Google  "busca recetas de paella"
  - Buscar YouTube    "encuentrame videos de gatos en YouTube"
  - Wikipedia         "que dice wikipedia sobre Newton"

SISTEMA
  - Abrir app         "abre la calculadora" / "abre notepad"
  - Listar carpeta    "que hay en mi escritorio"
  - Leer archivo      "leeme el fichero notas.txt"
  - Escribir archivo  "crea un fichero hola.txt con contenido..."
  - Ejecutar comando  "ejecuta ipconfig" (pide confirmacion)

MULTIMEDIA (todo se muestra EN el visor derecho)
  - Captura pantalla   "hazme una captura"
  - Generar imagen     "generame una imagen de un dragon volando"
  - Mini-video         "crea un video de 4 escenas sobre el espacio"
  - Resumir PDF        "resume el PDF ~/Downloads/libro.pdf"

DOCUMENTOS (via NOVA Office silencioso)
  - Word               "crea un word con un ensayo sobre la IA"
  - Excel              "hazme un excel con presupuesto mensual"

CONTROL COACH (via nova_coach)
  - Anadir habito      "anade el habito de leer 30 min"
  - Crear meta         "ponme la meta: correr 5K este mes"

INFO
  - Hora             "que hora es"
  - Clima            "como esta el clima en Sevilla"
  - Calcular         "cuanto es 234 * 891"

MEMORIA
  - Recordar hecho   "recuerda que mi cumple es 12 mayo"
  - Ver memoria      "que sabes de mi"
  - Olvidar          "olvida el hecho 3"

RECORDATORIOS
  - Crear            "recuerdame en 20 minutos comprar pan"
  - Listar           "que recordatorios tengo"

MODULOS NOVA
  - Abrir modulo     "abre NOVA Reader" / "lanza NOVA Studio"

COMANDOS RAPIDOS (con /)
  /nuevo         Nueva conversacion
  /memoria       Ver memoria
  /tools         Ver esta lista
  /ayuda         Ayuda general
  /recordar X    Recordar directamente el hecho X
  /limpiar       Limpiar chat"""

        txt = scrolledtext.ScrolledText(dlg, bg=C["bg2"], fg=C["fg"],
            font=("Consolas", 9), wrap="word", relief="flat",
            padx=14, pady=14)
        txt.pack(fill="both", expand=True, padx=12, pady=8)
        txt.insert("1.0", info)
        txt.config(state="disabled")

    def _ver_ayuda(self):
        ayuda = """NOVA JARVIS - Guia rapida

1) CONFIGURA la IA
   Boton 'Config IA' -> pon API key (Groq/Gemini/OpenAI...)
   La lista de modelos se carga automaticamente.

2) HABLA con Jarvis
   Escribe en el cuadro de abajo y ENVIAR.
   O activa 'Modo escucha continua' para conversar por voz.

3) PIDE LO QUE QUIERAS
   Jarvis tiene MUCHAS herramientas. Solo pide en lenguaje natural:
   - "que hora es"
   - "abre YouTube"
   - "recuerda que mi novia se llama Ana"
   - "generame una imagen de un robot amigable"
   - "recuerdame en 5 min tomar el cafe"

4) HERRAMIENTAS ON/OFF
   Si desactivas herramientas, Jarvis solo chatea (no ejecuta acciones).

5) VOZ
   - Voz ON/OFF: si Jarvis habla las respuestas
   - Escuchar una vez: activa micro una vez
   - Modo continuo: escucha-responde-escucha en bucle

6) MEMORIA
   Jarvis recuerda hechos importantes automaticamente.
   Boton 'Memoria' para ver/borrar.

DEPENDENCIAS OPCIONALES:
  pip install edge-tts pygame           (voz)
  pip install SpeechRecognition PyAudio (reconocimiento voz)
  pip install Pillow                    (capturas de pantalla)"""

        dlg = tk.Toplevel(self.root)
        dlg.title("Ayuda")
        dlg.geometry("640x600")
        dlg.configure(bg=C["bg"])
        txt = scrolledtext.ScrolledText(dlg, bg=C["bg2"], fg=C["fg"],
            font=("Consolas", 10), wrap="word", relief="flat",
            padx=14, pady=14)
        txt.pack(fill="both", expand=True, padx=12, pady=12)
        txt.insert("1.0", ayuda)
        txt.config(state="disabled")

    def _config(self):
        cfg_actual = {
            "proveedor": self.proveedor,
            "modelo": self.modelo,
            "api_key": self.api_key,
            "api_keys": self.cfg.get("api_keys", {}),
        }
        col = {"bg": C["bg"], "bg2": C["bg2"], "fg": C["fg"],
               "accent": C["jarvis"], "warn": C["warn"], "ok": C["ok"]}
        def on_save(n):
            self.proveedor = n["proveedor"]
            self.modelo = n["modelo"]
            self.api_key = n["api_key"]
            self.cfg["jarvis_proveedor"] = self.proveedor
            self.cfg["jarvis_modelo"] = self.modelo
            self.cfg["api_keys"] = n["api_keys"]
            guardar_json(CONFIG, self.cfg)
            self._chat_msg("Sistema",
                f"Nueva config: {self.proveedor} / {self.modelo}", "sys")
        dialogo_config(self.root, cfg_actual, on_save,
            titulo="Config IA - Jarvis", colores=col)

    # ============================================================
    # SISTEMA DE ACTUALIZACIONES
    # ============================================================
    def _comprobar_update_silencioso(self):
        """Comprueba updates al arrancar SIN mostrar dialogo si no hay nada nuevo."""
        if not UPDATER_OK:
            return
        def tarea():
            try:
                updater = Updater(MODULO_ID, VERSION)
                info = updater.comprobar()
                if info.get("disponible"):
                    # Mostrar banner en el chat
                    self.root.after(0, lambda: self._mostrar_banner_update(info))
            except Exception as e:
                print(f"[Update check] {e}")
        threading.Thread(target=tarea, daemon=True).start()

    def _mostrar_banner_update(self, info):
        """Muestra un mensaje discreto en el chat cuando hay update."""
        self._chat_msg("🆕 UPDATE",
            f"Version {info['version']} disponible (actual: {info['version_actual']}).\n"
            f"Pulsa 'Update' arriba para ver las novedades e instalar.",
            "warn")

    def _comprobar_update_manual(self):
        """Se dispara al pulsar el boton Update: muestra dialogo siempre."""
        if not UPDATER_OK:
            messagebox.showinfo("Sin updater",
                "El sistema de actualizaciones no esta disponible.\n"
                "Falta nova_updater.py en la carpeta.")
            return
        # Mostrar mensaje de "buscando..."
        self._chat_msg("Sistema",
            f"Buscando actualizaciones para {MODULO_ID} v{VERSION}...", "sys")

        def tarea():
            try:
                updater = Updater(MODULO_ID, VERSION)
                info = updater.comprobar(forzar=True)  # ignora cache
                if info.get("error"):
                    self.root.after(0, lambda:
                        messagebox.showerror("Error",
                            f"No se pudo consultar:\n{info['error']}"))
                    return
                if info.get("disponible"):
                    # Mostrar dialogo con changelog
                    self.root.after(0, lambda:
                        mostrar_dialogo_update(self.root, MODULO_ID, VERSION))
                else:
                    self.root.after(0, lambda:
                        messagebox.showinfo("Ya actualizado",
                            f"Ya tienes la ultima version: {VERSION} ✅"))
            except Exception as e:
                self.root.after(0, lambda e=e:
                    messagebox.showerror("Error", f"{e}"))
        threading.Thread(target=tarea, daemon=True).start()

    def _config_personalidad(self):
        """Dialogo para elegir personalidad."""
        if not PERS_OK:
            messagebox.showinfo("Sin personalidades",
                "Falta nova_personalidades.py en la carpeta.")
            return

        dlg = tk.Toplevel(self.root)
        dlg.title("🎭 Personalidad de Jarvis")
        dlg.geometry("520x600")
        dlg.configure(bg=C["bg"])
        dlg.transient(self.root)

        tk.Label(dlg, text="🎭 Elige la personalidad de Jarvis",
            bg=C["bg"], fg=C["jarvis"],
            font=("Segoe UI", 14, "bold")).pack(pady=12)

        tk.Label(dlg,
            text="La personalidad cambia el tono, el estilo y la voz de Jarvis.\n"
                 "Puedes cambiarla cuando quieras, y la conversacion sigue.",
            bg=C["bg"], fg=C["fg2"],
            font=("Segoe UI", 9, "italic")).pack(pady=4)

        pers_var = tk.StringVar(value=self.personalidad)

        # Frame con scroll
        cont = tk.Frame(dlg, bg=C["bg"])
        cont.pack(fill="both", expand=True, padx=20, pady=10)

        for pid, info in PERSONALIDADES.items():
            card = tk.Frame(cont, bg=C["bg2"], bd=1, relief="flat")
            card.pack(fill="x", pady=4)

            rb = tk.Radiobutton(card, text=info["nombre"],
                variable=pers_var, value=pid,
                bg=C["bg2"], fg=C["jarvis"], selectcolor=C["bg3"],
                activebackground=C["bg2"], activeforeground=C["jarvis"],
                font=("Segoe UI", 11, "bold"), anchor="w")
            rb.pack(fill="x", padx=10, pady=(6, 0))

            tk.Label(card, text=info["descripcion"],
                bg=C["bg2"], fg=C["fg2"],
                font=("Segoe UI", 9, "italic"),
                wraplength=440, justify="left", anchor="w"
                ).pack(fill="x", padx=32, pady=(0, 6))

        def guardar():
            self.personalidad = pers_var.get()
            self.cfg["jarvis_personalidad"] = self.personalidad
            guardar_json(CONFIG, self.cfg)
            # Cambiar voz preferida de esa personalidad
            if PERS_OK:
                nueva_voz = obtener_voz(self.personalidad)
                if nueva_voz:
                    self.voz = nueva_voz
                    self.cfg["jarvis_voz"] = nueva_voz
                    guardar_json(CONFIG, self.cfg)
            info = PERSONALIDADES.get(self.personalidad, {})
            self._chat_msg("Sistema",
                f"Personalidad cambiada a: {info.get('nombre', self.personalidad)}",
                "sys")
            dlg.destroy()

        tk.Button(dlg, text="💾 Aplicar personalidad", command=guardar,
            bg=C["ok"], fg="white", relief="flat",
            font=("Segoe UI", 11, "bold"),
            padx=22, pady=8, cursor="hand2").pack(pady=14)

    def _config_voz(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Voz de Jarvis")
        dlg.geometry("400x360")
        dlg.configure(bg=C["bg"])

        tk.Label(dlg, text="Elige la voz de Jarvis",
            bg=C["bg"], fg=C["jarvis"],
            font=("Segoe UI", 13, "bold")).pack(pady=12)

        voz_var = tk.StringVar()
        for label, codigo in VOCES.items():
            if codigo == self.voz:
                voz_var.set(label); break
        else:
            voz_var.set(list(VOCES.keys())[0])

        for label in VOCES:
            tk.Radiobutton(dlg, text=label, variable=voz_var, value=label,
                bg=C["bg"], fg=C["fg"], selectcolor=C["bg2"],
                activebackground=C["bg"], activeforeground=C["jarvis"],
                font=("Segoe UI", 10), anchor="w"
                ).pack(fill="x", padx=40, pady=2)

        def guardar():
            self.voz = VOCES[voz_var.get()]
            self.cfg["jarvis_voz"] = self.voz
            guardar_json(CONFIG, self.cfg)
            dlg.destroy()
            # Prueba
            if EDGE_OK and self.voz_activada:
                threading.Thread(target=lambda: hablar_voz(
                    "Hola, soy Jarvis. Estoy listo para ayudarte.",
                    self.voz), daemon=True).start()

        tk.Button(dlg, text="Guardar", command=guardar,
            bg=C["ok"], fg="white", relief="flat",
            font=("Segoe UI", 11, "bold"), padx=20, pady=6,
            cursor="hand2").pack(pady=14)


if __name__ == "__main__":
    # Splash screen bonita al arrancar (opcional, si falla se salta)
    # NOTA: si no funciona, comenta el bloque try/except entero
    import os as _os
    if _os.environ.get("NOVA_NO_SPLASH") != "1":
        try:
            from nova_ui import splash_screen, Theme
            splash_screen("NOVA Jarvis",
                           subtitulo="Asistente todopoderoso con IA",
                           color_acento=Theme.accent,
                           tareas=[
                               ("Cargando memoria...", None),
                               ("Preparando herramientas...", None),
                               ("Iniciando esfera cuantica...", None),
                           ],
                           duracion_min=1.0)
        except Exception as e:
            print(f"Splash no disponible (se ignora): {e}")

    print("Iniciando ventana principal de Jarvis...")
    try:
        root = tk.Tk()
        print("Ventana creada. Construyendo UI...")
        app = NovaJarvis(root)
        print("UI construida. Entrando en mainloop.")
        root.mainloop()
        print("Jarvis cerrado correctamente.")
    except Exception as e:
        import traceback
        print(f"\n\n[ERROR FATAL EN ARRANQUE]: {type(e).__name__}: {e}\n")
        traceback.print_exc()
        print("\nCopia este error completo y pasamelo.")
        input("\nPulsa ENTER para cerrar la ventana... ")
